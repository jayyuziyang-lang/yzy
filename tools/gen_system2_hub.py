#!/usr/bin/env python3
"""简易轮毂打紧机控制系统 - 梯形图+流程图+Word报告 一键生成"""

import os, xml.etree.ElementTree as ET

OUT = r'D:\Desktop\CAD1'
os.makedirs(OUT, exist_ok=True)

SYSTEM_NAME = '简易轮毂打紧机控制系统'
FILE_PREFIX = '简易轮毂打紧机控制系统'

IO_TABLE = [
    ['X000', '正转按钮 SB1',   'Y000', '电机正转 KM1',    'NO触点, 绿色'],
    ['X001', '反转按钮 SB2',   'Y001', '电机反转 KM2',    'NO触点, 绿色'],
    ['X002', '停止按钮 SB3',   'Y002', '正转指示灯 HL1',   '绿色'],
    ['X003', '急停按钮 SB4',   'Y003', '反转指示灯 HL2',   '黄色'],
    ['X004', '正转到位 SQ1',   'Y004', '流水灯1 HL4',     '1s循环'],
    ['X005', '反转到位 SQ2',   'Y005', '流水灯2 HL5',     '1s循环'],
    ['X006', '加速按钮 SB5',   'Y006', '流水灯3 HL6',     '1s循环'],
    ['X007', '减速按钮 SB6',   'Y007', '流水灯4 HL7',     '1s循环'],
    ['M0',   '正转使能(内部)',  'Y010', '低速 VFD-S1',     '变频器多段速1'],
    ['M1',   '反转使能(内部)',  'Y011', '中速 VFD-S2',     '变频器多段速2'],
    ['M2',   '系统运行中(内部)', 'Y012', '高速 VFD-S3',     '变频器多段速3'],
    ['C0',   '流水灯计数器',    'T0/T1', '闪烁定时器',     'K5=0.5s, 1s周期'],
    ['D0',   '速度等级 1-3',    '',     '',               ''],
]

LADDER_RUNGS = [
    {
        'label': '梯级1: 正转使能(自锁, 到位自动停, 反转互锁)',
        'elements': [
            ('NO', 80, 'X000\n正转键'),
            ('NC', 170, 'X004\n正转到位'),
            ('NC', 260, 'X002\n停止'),
            ('NC', 350, 'X003\n急停'),
            ('NC', 440, 'M1\n反转中'),
            ('COIL', 600, 'M0\n正转使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级2: 反转使能(自锁, 到位自动停, 正转互锁)',
        'elements': [
            ('NO', 80, 'X001\n反转键'),
            ('NC', 170, 'X005\n反转到位'),
            ('NC', 260, 'X002\n停止'),
            ('NC', 350, 'X003\n急停'),
            ('NC', 440, 'M0\n正转中'),
            ('COIL', 600, 'M1\n反转使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级3: 正转输出+反转输出(含互锁)',
        'elements': [
            ('NO', 80, 'M0\n正转使能'),
            ('NC', 170, 'Y001\n反转互锁'),
            ('COIL', 300, 'Y000\n正转'),
            ('COIL', 420, 'Y002\n正转灯'),
            ('NO', 520, 'M1\n反转使能'),
            ('NC', 610, 'Y000\n正转互锁'),
            ('COIL', 730, 'Y001\n反转'),
        ]
    },
    {
        'label': '梯级4: 反转指示灯',
        'elements': [
            ('NO', 80, 'Y001\n反转'),
            ('COIL', 200, 'Y003\n反转灯'),
        ]
    },
    {
        'label': '梯级5: 运行中标志(正转OR反转)',
        'elements': [
            ('NO', 80, 'M0\n正转使能'),
            ('COIL', 200, 'M2\n运行中'),
            ('NO', 300, 'M1\n反转使能'),
        ]
    },
    {
        'label': '梯级6: 1s周期闪烁(4流水灯循环)',
        'elements': [
            ('NO', 80, 'M2\n运行中'),
            ('NC', 160, 'T1'),
            ('COIL', 240, 'T0\nK5'),
            ('NO', 320, 'T0'),
            ('COIL', 400, 'T1\nK5'),
            ('NO', 480, 'T0'),
            ('FUNC', 560, 'INC\nC0'),
            ('FUNC', 640, '=\nC0 K4'),
            ('FUNC', 710, 'RST\nC0'),
        ]
    },
    {
        'label': '梯级7: 流水灯解码输出(Y4-Y7)',
        'elements': [
            ('FUNC', 80, '=\nC0 K0'),
            ('COIL', 180, 'Y004\n灯1'),
            ('FUNC', 270, '=\nC0 K1'),
            ('COIL', 370, 'Y005\n灯2'),
            ('FUNC', 460, '=\nC0 K2'),
            ('COIL', 560, 'Y006\n灯3'),
            ('FUNC', 650, '=\nC0 K3'),
            ('COIL', 750, 'Y007\n灯4'),
        ]
    },
    {
        'label': '梯级8: 速度调节(加速/减速)',
        'elements': [
            ('NO', 80, 'X006\n加速'),
            ('FUNC', 180, 'INCP\nD0'),
            ('NO', 290, 'X007\n减速'),
            ('FUNC', 390, 'DECP\nD0'),
        ]
    },
    {
        'label': '梯级9: 速度输出至变频器(D0=1低/2中/3高)',
        'elements': [
            ('FUNC', 80, '=\nD0 K1'),
            ('COIL', 180, 'Y010\n低速'),
            ('FUNC', 280, '=\nD0 K2'),
            ('COIL', 380, 'Y011\n中速'),
            ('FUNC', 480, '=\nD0 K3'),
            ('COIL', 580, 'Y012\n高速'),
        ]
    },
    {
        'label': '梯级10: 限位保护(无M0/M1时清除输出)',
        'elements': [
            ('NC', 80, 'M2\n未运行'),
            ('FUNC', 170, 'RST\nY000'),
            ('FUNC', 260, 'RST\nY001'),
            ('FUNC', 350, 'RST\nY004-Y007'),
        ]
    },
]

FLOW_NODES = [
    ('start', '开始\n(系统上电初始化)', 325, 30, 140, 45, 'startend', '#1A56DB'),
    ('init', '参数初始化\nD0=1(默认低速)\nC0=0', 325, 100, 140, 50, 'box', '#E8F0FE'),
    ('chk_estop', '急停按钮\n是否释放?', 325, 180, 110, 50, 'diamond', '#FFF3CD'),
    ('wait_cmd', '等待操作指令', 325, 260, 120, 45, 'box', '#E8F0FE'),
    ('chk_fwd', '正转键\n按下?', 200, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_rev', '反转键\n按下?', 450, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('fwd_run', 'M0=1 正转\nY000=1 正转输出\nY002=1 正转灯亮', 200, 430, 150, 55, 'box', '#D4EDDA'),
    ('rev_run', 'M1=1 反转\nY001=1 反转输出\nY003=1 反转灯亮', 450, 430, 150, 55, 'box', '#D4EDDA'),
    ('run_loop', '运行中 M2=1\n4流水灯循环(1s)', 325, 520, 160, 50, 'box', '#D4EDDA'),
    ('chk_pos', '到达指定\n位置?', 325, 600, 110, 50, 'diamond', '#FFF3CD'),
    ('auto_stop', '自动停止\nM0/M1=0\nY000/Y001=OFF', 325, 680, 150, 50, 'box', '#F8D7DA'),
    ('chk_stop', '停止按钮\n按下?', 150, 680, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_speed', '加速/减速\n按下?', 500, 680, 100, 50, 'diamond', '#FFF3CD'),
    ('adj_speed', '更新D0\n(INCP/DECP)', 500, 770, 120, 45, 'box', '#FCE4D6'),
    ('stop_act', '所有输出复位\nM0/M1=0', 150, 770, 130, 45, 'box', '#F8D7DA'),
    ('end', '结束', 150, 850, 100, 35, 'startend', '#6C757D'),
]

FLOW_CONNECTIONS = [
    ('start', 'init', 'down', ''),
    ('init', 'chk_estop', 'down', ''),
    ('chk_estop', 'chk_estop', 'right_loop', '急停'),
    ('chk_estop', 'wait_cmd', 'right', '是'),
    ('wait_cmd', 'chk_fwd', 'down', ''),
    ('wait_cmd', 'chk_rev', 'down', ''),
    ('chk_fwd', 'fwd_run', 'down', '是'),
    ('chk_fwd', 'wait_cmd', 'left_loop', '否'),
    ('chk_rev', 'rev_run', 'down', '是'),
    ('chk_rev', 'wait_cmd', 'right_loop', '否'),
    ('fwd_run', 'run_loop', 'right', ''),
    ('rev_run', 'run_loop', 'left', ''),
    ('run_loop', 'chk_pos', 'down', ''),
    ('run_loop', 'chk_stop', 'left', ''),
    ('run_loop', 'chk_speed', 'right', ''),
    ('chk_pos', 'auto_stop', 'down', '是'),
    ('chk_pos', 'run_loop', 'right_loop', '否'),
    ('auto_stop', 'wait_cmd', 'up_loop', ''),
    ('chk_stop', 'stop_act', 'down', '是'),
    ('chk_stop', 'run_loop', 'left_loop', '否'),
    ('stop_act', 'end', 'down', ''),
    ('chk_speed', 'adj_speed', 'down', '是'),
    ('chk_speed', 'run_loop', 'right_loop', '否'),
    ('adj_speed', 'run_loop', 'up_loop', ''),
    ('chk_estop', 'stop_act', 'left_estop', '急停'),
]


def make_ladder_svg(name, rungs, io_table):
    W = 850
    rung_h = 78
    left_rail, right_rail, coil_x = 45, 740, 580
    rung_start_y, contact_y_offset = 55, 35
    total_h = len(rungs) * rung_h + 60 + 280
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {total_h}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<style>text{font-size:10px;fill:#1a1a1a}.io-label{font-size:8px;fill:#333;text-anchor:middle}.rung-label{font-size:9px;fill:#666;text-anchor:middle}.contact-no{fill:none;stroke:#1a1a1a;stroke-width:1.5}.func-box{fill:#f5f5f5;stroke:#1a1a1a;stroke-width:1.5}.wire{stroke:#1a1a1a;stroke-width:1.5;fill:none}.rail{fill:#333}</style>',
             f'<rect width="{W}" height="{total_h}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — PLC梯形图</text>',
             f'<text x="{W/2}" y="44" text-anchor="middle" font-size="10" fill="#666">PLC: 三菱FX3U | 变频器多段速 | 4流水灯1s周期 | 到位自动停</text>']

    for i, rung in enumerate(rungs):
        y = rung_start_y + i * rung_h
        ey = y + contact_y_offset
        lines.append(f'<rect x="0" y="{y-3}" width="{W}" height="{rung_h-4}" fill="none" stroke="#e8e8e8" stroke-width="0.5"/>')
        lines.append(f'<rect x="{left_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<rect x="{right_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<line x1="{left_rail+3}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>')
        lines.append(f'<text x="{W/2}" y="{y+rung_h-6}" text-anchor="middle" class="rung-label">{rung["label"]}</text>')

        for elem in rung['elements']:
            typ, ex, lbl = elem
            if typ == 'NO':
                lines += [f'<line x1="{ex}" y1="{ey-10}" x2="{ex}" y2="{ey-16}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'NC':
                lines += [f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex+5}" y1="{ey-10}" x2="{ex-5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey-16}" x2="{ex}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'COIL':
                cx = ex
                lines += [f'<ellipse cx="{cx}" cy="{ey}" rx="13" ry="10" fill="none" stroke="#1a1a1a" stroke-width="1.5"/>',
                          f'<path d="M{cx-7},{ey-5} Q{cx-9},{ey} {cx-7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<path d="M{cx+7},{ey-5} Q{cx+9},{ey} {cx+7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<text x="{cx}" y="{ey+28}" class="io-label">{lbl}</text>',
                          f'<line x1="{cx+13}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>']
            elif typ == 'FUNC':
                fw, fh = 48, 28
                lines += [f'<rect x="{ex-fw/2}" y="{ey-fh/2}" width="{fw}" height="{fh}" rx="3" class="func-box"/>',
                          f'<text x="{ex}" y="{ey+4}" text-anchor="middle" font-size="9">{lbl}</text>']

        if rung.get('self_hold'):
            bx = rung['branch_from']
            sx = rung['self_x']
            by2 = ey + 46
            lines += [f'<line x1="{bx}" y1="{ey}" x2="{bx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{left_rail+6}" y1="{by2}" x2="{sx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{sx}" y1="{by2-12}" x2="{sx}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx}" y1="{by2+4}" x2="{sx}" y2="{by2+12}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2-4}" x2="{sx+5}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2+4}" x2="{sx+5}" y2="{by2+4}" class="contact-no"/>',
                      f'<text x="{sx}" y="{by2+26}" class="io-label">M{0 if '正转' in rung['label'] else '1'} 自锁</text>',
                      f'<line x1="{sx+10}" y1="{by2}" x2="{bx}" y2="{by2}" class="wire"/>']

    # I/O Table
    ty = rung_start_y + len(rungs) * rung_h + 12
    cw = [68, 100, 68, 100, 145]
    cx_pos = [25]
    for i in range(1, len(cw)): cx_pos.append(cx_pos[-1] + cw[i-1])
    lines.append(f'<text x="25" y="{ty}" font-size="13" font-weight="bold" fill="#1A56DB">I/O地址分配表 (三菱FX3U)</text>')
    tsy = ty + 12
    for j, (hdr, cx) in enumerate(zip(['输入端', '说明', '输出端', '说明', '备注'], cx_pos)):
        lines.append(f'<rect x="{cx}" y="{tsy}" width="{cw[j]}" height="20" fill="#1A56DB" stroke="#1A56DB"/>')
        lines.append(f'<text x="{cx+cw[j]/2}" y="{tsy+14}" text-anchor="middle" font-size="9" font-weight="bold" fill="white">{hdr}</text>')
    for ri, row in enumerate(io_table):
        ry = tsy + 20 + ri * 17
        fc = '#f8f9fa' if ri % 2 == 0 else '#fff'
        for j, (val, cx) in enumerate(zip(row, cx_pos)):
            lines.append(f'<rect x="{cx}" y="{ry}" width="{cw[j]}" height="17" fill="{fc}" stroke="#ddd" stroke-width="0.5"/>')
            lines.append(f'<text x="{cx+cw[j]/2}" y="{ry+12}" text-anchor="middle" font-size="8" fill="#333">{val}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_flow_svg(name, nodes, connections):
    W, H = 650, 960
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<defs><marker id="ah" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#333"/></marker></defs>',
             f'<rect width="{W}" height="{H}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — 软件流程图</text>',
             f'<text x="{W/2}" y="42" text-anchor="middle" font-size="10" fill="#666">主程序循环扫描 | 变频器三速 | 4流水灯1s周期 | 到位自动停</text>']

    nc = {}
    for n in nodes:
        nid, txt, cx, cy, w, h, shape, color = n
        nc[nid] = (cx, cy, w, h, shape)
        x, y = cx - w//2, cy - h//2
        if shape == 'startend':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="20" ry="20" fill="{color}" stroke="{color}"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="white" font-weight="bold">{ln}</text>')
        elif shape == 'diamond':
            lines.append(f'<polygon points="{cx},{y} {x+w},{cy} {cx},{y+h} {x},{cy}" fill="{color}" stroke="#856404" stroke-width="1.5"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#856404">{ln}</text>')
        elif shape == 'box':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="4" fill="{color}" stroke="#1A56DB" stroke-width="1.2"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#1a1a1a">{ln}</text>')

    for conn in connections:
        fid, tid, direction, label = conn
        fx, fy, fw, fh, _ = nc[fid]
        tx, ty, tw, th, _ = nc[tid]
        if direction == 'down':
            y1, y2 = fy + fh//2, ty - th//2
            lines.append(f'<line x1="{fx}" y1="{y1}" x2="{tx}" y2="{y2}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{fx+16}" y="{(y1+y2)/2+4}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'right':
            x1, x2 = fx + fw//2, tx - tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left_loop':
            xs = fx - fw//2
            pts = f'{xs},{fy} {xs-25},{fy} {xs-25},{ty} {tx+tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs-28}" y="{fy-6}" font-size="9" fill="#DC2626">{label}</text>')
        elif direction == 'right_loop':
            xs = fx + fw//2
            pts = f'{xs},{fy} {xs+28},{fy} {xs+28},{ty} {tx-tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs+30}" y="{ty-8}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'up_loop':
            xs = fx - fw//2
            pts = f'{xs},{fy} {xs-40},{fy} {xs-40},{ty} {tx},{ty+th//2}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
        elif direction == 'left_estop':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#DC2626" stroke-width="2" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-8}" text-anchor="middle" font-size="10" fill="#DC2626" font-weight="bold">{label}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_docx(name, prefix, io_table):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = '宋体'

    for _ in range(4):
        doc.add_paragraph('')

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('沈阳城市建设学院'); run.font.size = Pt(22); run.font.bold = True; run.font.name = '黑体'
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('实训课程任务书'); run.font.size = Pt(20); run.font.bold = True; run.font.name = '黑体'
    doc.add_paragraph('')

    for label, value in [('实训课程名称', '电气控制实训'), ('实训课程题目', name), ('学期', '2026年春季学期'),
                          ('学    院', '信息与控制工程学院'), ('专    业', '自动化（专升本）'),
                          ('班    级', '25级 2班'), ('指导教师', '王丹阳    职称   讲师')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}'); run.font.size = Pt(14); run.font.name = '宋体'
    doc.add_page_break()

    # 一、设计任务与要求
    h = doc.add_heading('一、设计任务与要求', level=1)
    tasks = [
        '设计一款基于可编程控制器（PLC）的简易轮毂打紧机控制系统，运用PLC、变频器、按钮指示灯等设备',
        '实现电机的启动、停止、正转、反转功能，速度可调',
        '按正转键，电机正转，正转指示灯点亮，到达指定位置（限位开关）自动停止',
        '按反转键，电机反转，反转指示灯点亮，到达指定位置（限位开关）自动停止',
        '4个流水灯循环点亮，1秒周期',
        '按下停止键，电机停止转动，所有指示灯熄灭',
        '系统考虑安全信号，异常情况急停按钮停止系统运行',
        '通过触摸屏显示运行/停止状态',
        '具备电气画图能力和PLC编程设计能力',
    ]
    for t in tasks: doc.add_paragraph(t, style='List Bullet')

    # 二、器件选型
    doc.add_heading('二、器件选型', level=1)
    components = [
        ['可编程控制器PLC', '三菱 FX3U-32MR/ES-A', '1台', '16入/16出，继电器输出'],
        ['变频器', '三菱 FR-D720S-0.75K-CHT', '1台', '0.75kW，多段速控制'],
        ['三相异步电机', 'YS-7124', '1台', '0.37kW，380V，1400rpm'],
        ['断路器', 'DZ47-63 3P C10', '1个', '主电路短路保护'],
        ['熔断器', 'RT18-32 6A', '2个', '控制电路过载保护'],
        ['交流接触器', 'CJX2-0910 220V', '2个', '正反转控制，机械互锁'],
        ['热继电器', 'JR36-20 1-1.6A', '1个', '电机过载保护'],
        ['正转按钮', 'LAY39-11BN 绿色', '1个', '自复位常开触点'],
        ['反转按钮', 'LAY39-11BN 黄色', '1个', '自复位常开触点'],
        ['停止按钮', 'LAY39-11BN 红色', '1个', '自复位常闭触点'],
        ['急停按钮', 'LAY39-11ZS 红色', '1个', '旋转复位，常闭触点'],
        ['限位开关(正转到位)', 'LX19-001', '1个', '滚轮式，常闭触点'],
        ['限位开关(反转到位)', 'LX19-001', '1个', '滚轮式，常闭触点'],
        ['流水指示灯', 'AD16-22DS 彩色', '4个', '红黄绿蓝各1，1s循环'],
        ['正转指示灯', 'AD16-22DS 绿色', '1个', '正转状态指示'],
        ['反转指示灯', 'AD16-22DS 黄色', '1个', '反转状态指示'],
        ['触摸屏', '威纶通 TK6071iQ', '1台', '7寸，RS485通信'],
        ['开关电源', 'S-100-24', '1个', 'DC24V 4.5A'],
        ['导线及辅材', 'BVR 1.5mm² 等', '若干', '控制线及动力线'],
    ]
    table = doc.add_table(rows=len(components)+1, cols=5); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['序号', '器件名称', '型号规格', '数量', '备注']):
        cell = table.rows[0].cells[j]; cell.text = hdr
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
    for i, row_data in enumerate(components):
        table.rows[i+1].cells[0].text = str(i+1)
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j+1].text = val
            for p in table.rows[i+1].cells[j+1].paragraphs:
                for run in p.runs: run.font.size = Pt(9)

    # 三、IO表
    doc.add_heading('三、I/O地址分配', level=1)
    io_table_obj = doc.add_table(rows=len(io_table)+1, cols=5); io_table_obj.style = 'Table Grid'; io_table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['输入端', '说明', '输出端', '说明', '备注']):
        cell = io_table_obj.rows[0].cells[j]; cell.text = hdr
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.bold = True; run.font.size = Pt(9)
    for i, row_data in enumerate(io_table):
        for j, val in enumerate(row_data):
            io_table_obj.rows[i+1].cells[j].text = val
            for p in io_table_obj.rows[i+1].cells[j].paragraphs:
                for run in p.runs: run.font.size = Pt(8)

    # 四、电路分析
    doc.add_heading('四、电路分析', level=1)
    doc.add_heading('4.1 主电路分析', level=2)
    doc.add_paragraph('主电路由断路器QF引入三相380V电源，经熔断器FU1保护后，通过交流接触器KM1（正转/打紧）和KM2（反转/松开）控制三相异步电机。KM1与KM2之间设置机械互锁和电气互锁，防止相间短路。热继电器FR提供电机过载保护。变频器VFD输入端经接触器连接工频电源，输出端连接电机，通过多段速端子（S1-S3）选择不同运行频率，以适应不同规格轮毂的打紧力矩需求。')
    doc.add_paragraph('限位开关SQ1和SQ2分别安装在正转和反转方向的指定位置，当工作台到达设定位置时，限位开关动作，通过PLC程序自动停止相应方向的运行。')
    doc.add_heading('4.2 控制电路分析', level=2)
    doc.add_paragraph('控制电路采用DC24V供电。PLC输入端X000~X007分别接收正转键、反转键、停止键、急停键、正转限位、反转限位和加速/减速按钮信号。输出端Y000~Y007和Y010~Y012分别控制正反转接触器、指示灯和变频器多段速端子。4个流水灯（Y004~Y007）通过计数器C0和比较指令实现1秒周期的循环点亮效果。限位开关X004/X005的常闭触点串联在正转/反转使能回路中，到达位置后自动断开，实现自动停机。')

    # 五、梯形图
    doc.add_heading('五、梯形图程序设计', level=1)
    ladder_exp = [
        ('梯级1-2：正转/反转使能（自锁+到位自动停+互锁）',
         '按下X000（正转键），M0得电自锁，正转运行。串联的常闭触点X004（正转限位）到达指定位置后断开→M0失电自动停止。反转同理，X005检测反转到位。M0与M1之间通过常闭触点互锁。X002（停止）和X003（急停）用于紧急情况下强制停止。'),
        ('梯级3-4：正反转输出及指示灯',
         'M0→Y000正转输出+Y002正转灯亮；M1→Y001反转输出+Y003反转灯亮。正反转输出回路中串联对方输出继电器的常闭触点（Y000/Y001），构成软件互锁。'),
        ('梯级5：运行中标志', 'M0或M1任一接通→M2=ON，M2作为流水灯及其他共用的运行标志。'),
        ('梯级6：1s周期闪烁+流水灯计数',
         'M2运行→T0/T1交替振荡（各K5=0.5s，1s周期）。T0每0.5s产生一个上升沿→计数器C0加1（INC C0）。当C0=4时，RST C0将计数器清零→C0值在0→1→2→3→0之间循环。'),
        ('梯级7：4流水灯解码输出',
         'C0=0→Y004亮；C0=1→Y005亮；C0=2→Y006亮；C0=3→Y007亮。随着C0以1s周期变化，4个流水灯依次点亮，形成流水效果。'),
        ('梯级8-9：速度调节与输出',
         'X006（加速）→INCP D0（D0+1，上限3）；X007（减速）→DECP D0（D0-1，下限1）。D0值比较后输出到Y010/Y011/Y012（变频器多段速端子），实现三档调速。'),
        ('梯级10：安全保护', 'M2未运行时（系统停止），批量复位Y000/Y001/Y004~Y007，确保所有输出可靠关闭。'),
    ]
    for title, exp in ladder_exp:
        doc.add_heading(title, level=2); doc.add_paragraph(exp)
    doc.add_paragraph(f'完整梯形图详见附件：{prefix}梯形图.svg')

    # 六、流程图
    doc.add_heading('六、软件流程图', level=1)
    doc.add_paragraph('系统上电初始化后进入主循环。首先检测急停状态，若急停按下则直接停止。然后扫描正转键和反转键，响应后启动对应方向的运行（含自锁和互锁）。运行过程中同时检测限位开关（到位自动停）、停止按钮、加速/减速按钮。4个流水灯在运行期间持续以1s周期循环。完整流程详见附件。')
    doc.add_paragraph(f'完整软件流程图详见附件：{prefix}软件流程图.svg')

    # 七、功能说明
    doc.add_heading('七、系统功能说明', level=1)
    for fd in [
        '正转打紧：按下正转键（X000）→电机正转（Y000），正转指示灯（Y002）亮，工作台前进进行轮毂打紧。到达设定位置后正转限位（X004）断开，自动停止。',
        '反转松开：按下反转键（X001）→电机反转（Y001），反转指示灯（Y003）亮，工作台后退。到达设定位置后反转限位（X005）断开，自动停止。',
        '速度调节：通过加速（X006）/减速（X007）按钮无极切换三档速度（D0=1/2/3），适应不同规格轮毂的力矩需求。变频器根据D0值通过多段速端子（Y010/Y011/Y012）输出对应频率。',
        '4流水灯：系统运行中（M2=ON），4个流水灯（Y4→Y5→Y6→Y7→Y4）以1s周期循环点亮，直观显示运行状态。',
        '安全保护：正反转互锁（软件+硬件双重保护）；急停按钮（X003）最高优先级；限位开关自动停机；未运行时所有输出强制复位。',
    ]: doc.add_paragraph(fd, style='List Bullet')

    # 八、心得体会
    doc.add_heading('八、心得体会', level=1)
    for t in [
        f'通过本次《{name}》实训课程设计，我对PLC在工业控制领域的应用有了更深入的理解。',
        '一、对轮毂打紧工艺的理解：轮毂打紧是汽车装配线中的重要工序，需要精确控制力矩和位置。PLC与变频器的配合使用，可以灵活调节电机转速（对应打紧力矩），并通过限位开关实现精确的位置控制。这种"速度可调+位置控制"的方案在工业自动化中具有广泛的适用性。',
        '二、PLC流水灯编程的收获：流水灯是PLC入门经典案例，但4位流水灯配合1s周期的计数器实现，需要掌握定时器级联（T0→T1振荡）、计数器循环（C0 0-3循环）和比较输出（CMP指令解码）三个核心概念。本次设计让我对三菱FX系列的计数器、比较指令有了熟练的掌握。',
        '三、位置检测与自动停止：限位开关作为反馈元件，将物理位置转化为电信号，PLC通过常闭触点串联在运行回路中实现自动停止。这种"传感器→PLC→执行器"的闭环逻辑是工业自动化的基本范式。',
        '四、安全设计的系统性思考：本次设计从硬件（急停按钮直接切断、限位开关机械检测、接触器机械互锁）和软件（梯形图电气互锁、输出强制复位、自锁解除）两个维度构建了完整的安全体系。安全设计不能只依赖单一层面，必须做到纵深防御（Defense in Depth）。',
        '五、总结与展望：本次实训完整经历了一个PLC控制系统从需求分析到程序设计的全过程。虽然只是一个简化版的轮毂打紧机模型，但其核心控制逻辑（正反转+互锁+位置检测+速度调节+流水指示）与实际工业系统完全一致。这为将来从事自动化相关工作打下了坚实基础。',
    ]: doc.add_paragraph(t)

    # 九、参考文献
    doc.add_heading('九、参考文献', level=1)
    for ref in ['三菱电机. FX3U系列微型可编程控制器用户手册[硬件篇]. 2018.',
                 '三菱电机. FX系列微型可编程控制器编程手册[基本指令篇]. 2019.',
                 '三菱电机. FR-D720S系列变频器使用手册. 2020.',
                 '廖常初. FX系列PLC编程及应用(第4版). 机械工业出版社, 2021.',
                 '王曙光. 电气控制与PLC应用技术. 清华大学出版社, 2020.',
                 '陈建明. 电气控制与PLC应用(第3版). 电子工业出版社, 2019.',
                 '威纶通科技. TK6071iQ触摸屏用户手册. 2021.',
                 'GB/T 15969.3-2017 可编程序控制器 第3部分: 编程语言[S]. 中国标准出版社, 2017.']:
        doc.add_paragraph(ref)

    docx_path = os.path.join(OUT, f'{prefix}.docx')
    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    print(f'=== 生成 {SYSTEM_NAME} ===')
    ladder_svg = make_ladder_svg(SYSTEM_NAME, LADDER_RUNGS, IO_TABLE)
    ladder_path = os.path.join(OUT, f'{FILE_PREFIX}梯形图.svg')
    with open(ladder_path, 'w', encoding='utf-8') as f: f.write(ladder_svg)
    print(f'[OK] 梯形图: {ladder_path} ({len(ladder_svg)} bytes)')

    flow_svg = make_flow_svg(SYSTEM_NAME, FLOW_NODES, FLOW_CONNECTIONS)
    flow_path = os.path.join(OUT, f'{FILE_PREFIX}软件流程图.svg')
    with open(flow_path, 'w', encoding='utf-8') as f: f.write(flow_svg)
    print(f'[OK] 流程图: {flow_path} ({len(flow_svg)} bytes)')

    docx_path = make_docx(SYSTEM_NAME, FILE_PREFIX, IO_TABLE)
    print(f'[OK] Word报告: {docx_path}')

    for fp, fn in [(ladder_path, '梯形图'), (flow_path, '流程图')]:
        try:
            ET.fromstring(open(fp, 'r', encoding='utf-8').read())
            print(f'[OK] {fn} XML语法通过')
        except ET.ParseError as e: print(f'[FAIL] {fn} XML: {e}')
    print('完成!')
