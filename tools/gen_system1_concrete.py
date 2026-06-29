#!/usr/bin/env python3
"""混凝土搅拌控制系统 - 梯形图+流程图+Word报告 一键生成"""

import os, sys, xml.etree.ElementTree as ET

OUT = r'D:\Desktop\CAD1'
os.makedirs(OUT, exist_ok=True)

# ============================================================
# 系统参数
# ============================================================
SYSTEM_NAME = '混凝土搅拌控制系统'
FILE_PREFIX = '混凝土搅拌控制系统'
SPEEDS = ['15Hz', '20Hz', '30Hz']
FORWARD_TIME = 6  # 正转6秒
REVERSE_TIME = 6  # 反转6秒
STOP_DELAY = 3    # 停止延时3秒

# I/O分配
IO_TABLE = [
    ['X000', '启动按钮 SB1',   'Y000', '电机正转 KM1',   'NO触点, 绿色'],
    ['X001', '停止按钮 SB2',   'Y001', '电机反转 KM2',   'NC触点, 红色'],
    ['X002', '急停按钮 SB3',   'Y002', '运行指示灯 HL1',  '绿色, 运行时常亮'],
    ['X003', '15Hz按钮 SB4',   'Y003', '正转指示灯 HL2',  '绿色'],
    ['X004', '20Hz按钮 SB5',   'Y004', '反转指示灯 HL3',  '黄色'],
    ['X005', '30Hz按钮 SB6',   'Y010', '15Hz输出 VFD-S1', '变频器多段速1'],
    ['M0',   '运行使能(内部)',  'Y011', '20Hz输出 VFD-S2', '变频器多段速2'],
    ['M1',   '方向记忆(内部)',  'Y012', '30Hz输出 VFD-S3', '变频器多段速3'],
    ['M10',  '停止延时中(内部)', 'T0',  '正转定时器',       'K60 = 6秒'],
    ['D0',   '当前速度等级',     'T1',  '反转定时器',       'K60 = 6秒'],
    ['',     '',                'T2',  '停止延时定时器',    'K30 = 3秒'],
]

# 梯形图梯级定义
LADDER_RUNGS = [
    {
        'label': '梯级1: 系统运行使能(自锁, X001停止为NC, X002急停为NC)',
        'elements': [
            ('NO', 80, 'X000\n启动'),
            ('NC', 260, 'X001\n停止'),
            ('NC', 350, 'X002\n急停'),
            ('COIL', 600, 'M0\n运行使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级2: 正转6秒定时(运行时T0计时, 常闭T1复位)',
        'elements': [
            ('NO', 80, 'M0\n运行'),
            ('NC', 170, 'T1\n反转中'),
            ('COIL', 260, 'T0\nK60'),
            ('NC', 350, 'T0\n6s到'),
            ('COIL', 440, 'M1\n正转中'),
        ]
    },
    {
        'label': '梯级3: 反转6秒定时(T0计时到→T1开始计时)',
        'elements': [
            ('NO', 80, 'T0\n6s到'),
            ('NC', 170, 'T2\n停止延时'),
            ('COIL', 260, 'T1\nK60'),
        ]
    },
    {
        'label': '梯级4: 正转输出(含互锁, T0未到时正转)',
        'elements': [
            ('NO', 80, 'M0\n运行'),
            ('NC', 170, 'T0\n正转结束'),
            ('NC', 260, 'Y001\n反转互锁'),
            ('NC', 350, 'T2\n停止延时'),
            ('COIL', 600, 'Y000\n正转'),
        ]
    },
    {
        'label': '梯级5: 反转输出(含互锁, T0到-T1到时反转)',
        'elements': [
            ('NO', 80, 'T0\n正转结束'),
            ('NC', 170, 'T1'),
            ('NC', 260, 'Y000\n正转互锁'),
            ('NC', 350, 'T2\n停止延时'),
            ('COIL', 600, 'Y001\n反转'),
        ]
    },
    {
        'label': '梯级6: 停止延时3秒(X001停止→T2延时3s→M0复位)',
        'elements': [
            ('NO', 80, 'X001\n停止'),
            ('NC', 160, 'T2'),
            ('COIL', 240, 'T2\nK30'),
            ('NO', 320, 'T2\n3s到'),
            ('COIL', 420, 'M10\n延时到'),
        ]
    },
    {
        'label': '梯级7: 速度选择(自锁按钮, 互锁选择15/20/30Hz)',
        'elements': [
            ('NO', 80, 'X003\n15Hz'),
            ('COIL', 180, 'M3\n选15'),
            ('NO', 260, 'X004\n20Hz'),
            ('NC', 340, 'M3'),
            ('NC', 400, 'M5'),
            ('COIL', 480, 'M4\n选20'),
            ('NO', 560, 'X005\n30Hz'),
            ('NC', 630, 'M3'),
            ('NC', 680, 'M4'),
            ('COIL', 760, 'M5\n选30'),
        ]
    },
    {
        'label': '梯级8: 速度输出至变频器',
        'elements': [
            ('NO', 80, 'M3\n15Hz'),
            ('COIL', 200, 'Y010\n低速'),
            ('NO', 310, 'M4\n20Hz'),
            ('COIL', 430, 'Y011\n中速'),
            ('NO', 540, 'M5\n30Hz'),
            ('COIL', 660, 'Y012\n高速'),
        ]
    },
    {
        'label': '梯级9: 状态指示灯(正转=Y003, 反转=Y004, 运行=Y002)',
        'elements': [
            ('NO', 80, 'Y000\n正转'),
            ('COIL', 200, 'Y003\n正转灯'),
            ('NO', 310, 'Y001\n反转'),
            ('COIL', 430, 'Y004\n反转灯'),
            ('NO', 540, 'M0\n运行'),
            ('COIL', 660, 'Y002\n运行灯'),
        ]
    },
]

# 流程图节点
FLOW_NODES = [
    ('start', '开始\n(系统上电初始化)', 325, 30, 140, 45, 'startend', '#1A56DB'),
    ('init', '参数初始化\n默认15Hz\nT0/T1/T2清零', 325, 100, 150, 55, 'box', '#E8F0FE'),
    ('chk_estop', '急停按钮\n是否释放?', 325, 185, 130, 50, 'diamond', '#FFF3CD'),
    ('wait_start', '等待启动按钮', 325, 270, 140, 50, 'box', '#E8F0FE'),
    ('chk_start', '启动按钮\n按下?', 325, 350, 110, 50, 'diamond', '#FFF3CD'),
    ('set_run', 'M0=1 运行\n电机正转 Y000=1\nT0开始计时6s', 325, 430, 170, 55, 'box', '#D4EDDA'),
    ('fwd_loop', '正转中...\nT0计时中', 325, 515, 130, 50, 'box', '#E8F0FE'),
    ('chk_t0', 'T0=6s\n正转时间到?', 325, 595, 120, 50, 'diamond', '#FFF3CD'),
    ('rev_start', 'Y000=0(正转停)\nY001=1(反转启动)\nT1开始计时6s', 500, 600, 170, 55, 'box', '#D4EDDA'),
    ('rev_loop', '反转中...\nT1计时中', 500, 685, 130, 50, 'box', '#E8F0FE'),
    ('chk_t1', 'T1=6s\n反转时间到?', 500, 765, 120, 50, 'diamond', '#FFF3CD'),
    ('chk_stop', '停止按钮\n按下?', 325, 680, 110, 50, 'diamond', '#FFF3CD'),
    ('stop_dly', '停止延时3s\nT2=K30计时', 150, 685, 140, 50, 'box', '#FFF3CD'),
    ('chk_t2', 'T2=3s\n延时到?', 150, 765, 110, 50, 'diamond', '#F8D7DA'),
    ('stop_act', 'M0=0 停止\n所有输出复位\nY000~Y012=OFF', 150, 845, 160, 55, 'box', '#F8D7DA'),
    ('end', '结束', 150, 930, 100, 35, 'startend', '#6C757D'),
    ('chk_speed', '速度按钮\n按下?', 500, 850, 110, 50, 'diamond', '#FFF3CD'),
    ('set_speed', '更新速度等级\nD0=1/2/3', 500, 930, 130, 45, 'box', '#FCE4D6'),
]

FLOW_CONNECTIONS = [
    ('start', 'init', 'down', ''),
    ('init', 'chk_estop', 'down', ''),
    ('chk_estop', 'wait_start', 'right', '是'),
    ('wait_start', 'chk_start', 'down', ''),
    ('chk_start', 'set_run', 'right', '是'),
    ('chk_start', 'wait_start', 'left_loop', '否'),
    ('chk_estop', 'stop_act', 'left_estop', '急停'),
    ('set_run', 'fwd_loop', 'down', ''),
    ('fwd_loop', 'chk_stop', 'right', ''),
    ('fwd_loop', 'chk_t0', 'down', ''),
    ('chk_t0', 'rev_start', 'right', '是'),
    ('chk_t0', 'fwd_loop', 'left_loop', '否'),
    ('rev_start', 'rev_loop', 'down', ''),
    ('rev_loop', 'chk_t1', 'right', ''),
    ('rev_loop', 'chk_stop', 'left', ''),
    ('chk_t1', 'set_run', 'right_loop', '是(循环)'),
    ('chk_t1', 'rev_loop', 'left_loop', '否'),
    ('chk_stop', 'stop_dly', 'left', '是'),
    ('chk_stop', 'fwd_loop', 'right_loop', '否'),
    ('stop_dly', 'chk_t2', 'down', ''),
    ('chk_t2', 'stop_act', 'left', '是'),
    ('chk_t2', 'stop_dly', 'left_loop', '否'),
    ('stop_act', 'end', 'down', ''),
    ('rev_loop', 'chk_speed', 'down', ''),
    ('chk_speed', 'set_speed', 'right', '是'),
    ('chk_speed', 'chk_t1', 'right_loop_2', '否'),
    ('set_speed', 'rev_loop', 'up_loop', ''),
]


# ============================================================
# SVG 生成器
# ============================================================
def make_ladder_svg(name, rungs, io_table):
    W = 850
    rung_h = 80
    left_rail, right_rail, coil_x = 45, 740, 580
    rung_start_y, contact_y_offset = 55, 35

    total_h = len(rungs) * rung_h + 60 + 280
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {total_h}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<style>text{font-size:10px;fill:#1a1a1a}.io-label{font-size:8px;fill:#333;text-anchor:middle}.rung-label{font-size:9px;fill:#666;text-anchor:middle}.contact-no{fill:none;stroke:#1a1a1a;stroke-width:1.5}.func-box{fill:#f5f5f5;stroke:#1a1a1a;stroke-width:1.5}.wire{stroke:#1a1a1a;stroke-width:1.5;fill:none}.rail{fill:#333}</style>',
             f'<rect width="{W}" height="{total_h}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — PLC梯形图</text>',
             f'<text x="{W/2}" y="44" text-anchor="middle" font-size="10" fill="#666">PLC: 三菱FX3U | 变频器多段速 | 正反转自动循环 | 停止延时3s</text>']

    for i, rung in enumerate(rungs):
        y = rung_start_y + i * rung_h
        ey = y + contact_y_offset
        lines.append(f'<rect x="0" y="{y-3}" width="{W}" height="{rung_h-4}" fill="none" stroke="#e8e8e8" stroke-width="0.5"/>')
        lines.append(f'<rect x="{left_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<rect x="{right_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<line x1="{left_rail+3}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>')
        lines.append(f'<text x="{W/2}" y="{y+rung_h-6}" text-anchor="middle" class="rung-label">{rung["label"]}</text>')

        for elem in rung['elements']:
            typ, ex, lbl = elem
            if typ == 'NO':
                lines += [f'<line x1="{ex}" y1="{ey-10}" x2="{ex}" y2="{ey-16}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'NC':
                lines += [f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex+5}" y1="{ey-10}" x2="{ex-5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey-16}" x2="{ex}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'COIL':
                cx = ex
                lines += [f'<ellipse cx="{cx}" cy="{ey}" rx="13" ry="10" fill="none" stroke="#1a1a1a" stroke-width="1.5"/>',
                          f'<path d="M{cx-7},{ey-5} Q{cx-9},{ey} {cx-7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<path d="M{cx+7},{ey-5} Q{cx+9},{ey} {cx+7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<text x="{cx}" y="{ey+28}" class="io-label">{lbl}</text>',
                          f'<line x1="{cx+13}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>']
            elif typ == 'FUNC':
                fw, fh = 48, 28
                lines += [f'<rect x="{ex-fw/2}" y="{ey-fh/2}" width="{fw}" height="{fh}" rx="3" class="func-box"/>',
                          f'<text x="{ex}" y="{ey+4}" text-anchor="middle" font-size="9">{lbl}</text>']

        # Self-holding branch
        if rung.get('self_hold'):
            bx = rung['branch_from']
            sx = rung['self_x']
            by2 = ey + 48
            lines += [f'<line x1="{bx}" y1="{ey}" x2="{bx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{left_rail+6}" y1="{by2}" x2="{sx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{sx}" y1="{by2-12}" x2="{sx}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx}" y1="{by2+4}" x2="{sx}" y2="{by2+12}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2-4}" x2="{sx+5}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2+4}" x2="{sx+5}" y2="{by2+4}" class="contact-no"/>',
                      f'<text x="{sx}" y="{by2+26}" class="io-label">M0 自锁</text>',
                      f'<line x1="{sx+10}" y1="{by2}" x2="{bx}" y2="{by2}" class="wire"/>']

    # I/O Table
    ty = rung_start_y + len(rungs) * rung_h + 12
    cw = [68, 100, 68, 100, 145]
    cx_pos = [25]
    for i in range(1, len(cw)): cx_pos.append(cx_pos[-1] + cw[i-1])
    lines.append(f'<text x="25" y="{ty}" font-size="13" font-weight="bold" fill="#1A56DB">I/O地址分配表 (三菱FX3U)</text>')
    tsy = ty + 12
    for j, (hdr, cx) in enumerate(zip(['输入端', '说明', '输出端', '说明', '备注'], cx_pos)):
        lines.append(f'<rect x="{cx}" y="{tsy}" width="{cw[j]}" height="20" fill="#1A56DB" stroke="#1A56DB"/>')
        lines.append(f'<text x="{cx+cw[j]/2}" y="{tsy+14}" text-anchor="middle" font-size="9" font-weight="bold" fill="white">{hdr}</text>')
    for ri, row in enumerate(io_table):
        ry = tsy + 20 + ri * 17
        fc = '#f8f9fa' if ri % 2 == 0 else '#fff'
        for j, (val, cx) in enumerate(zip(row, cx_pos)):
            lines.append(f'<rect x="{cx}" y="{ry}" width="{cw[j]}" height="17" fill="{fc}" stroke="#ddd" stroke-width="0.5"/>')
            lines.append(f'<text x="{cx+cw[j]/2}" y="{ry+12}" text-anchor="middle" font-size="8" fill="#333">{val}</text>')

    lines.append('</svg>')
    return '\n'.join(lines)


def make_flow_svg(name, nodes, connections):
    W, H = 650, 1050
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<defs><marker id="ah" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#333"/></marker></defs>',
             f'<rect width="{W}" height="{H}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — 软件流程图</text>',
             f'<text x="{W/2}" y="42" text-anchor="middle" font-size="10" fill="#666">主程序循环扫描 | 变频器三速 | 正转6s⇄反转6s | 停止延时3s</text>']

    nc = {}
    for n in nodes:
        nid, txt, cx, cy, w, h, shape, color = n
        nc[nid] = (cx, cy, w, h, shape)
        x, y = cx - w//2, cy - h//2
        if shape == 'startend':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="20" ry="20" fill="{color}" stroke="{color}"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="white" font-weight="bold">{ln}</text>')
        elif shape == 'diamond':
            lines.append(f'<polygon points="{cx},{y} {x+w},{cy} {cx},{y+h} {x},{cy}" fill="{color}" stroke="#856404" stroke-width="1.5"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#856404">{ln}</text>')
        elif shape == 'box':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="4" fill="{color}" stroke="#1A56DB" stroke-width="1.2"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#1a1a1a">{ln}</text>')

    for conn in connections:
        fid, tid, direction, label = conn
        fx, fy, fw, fh, _ = nc[fid]
        tx, ty, tw, th, _ = nc[tid]
        if direction == 'down':
            y1, y2 = fy + fh//2, ty - th//2
            lines.append(f'<line x1="{fx}" y1="{y1}" x2="{tx}" y2="{y2}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{fx+16}" y="{(y1+y2)/2+4}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'right':
            x1, x2 = fx + fw//2, tx - tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left_loop':
            xs = fx - fw//2
            pts = f'{xs},{fy} {xs-25},{fy} {xs-25},{ty} {tx+tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs-28}" y="{fy-6}" font-size="9" fill="#DC2626">{label}</text>')
        elif direction == 'right_loop':
            xs = fx + fw//2
            pts = f'{xs},{fy} {xs+28},{fy} {xs+28},{ty} {tx-tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs+30}" y="{ty-8}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'right_loop_2':
            xs = fx + fw//2
            pts = f'{xs},{fy} {xs+40},{fy} {xs+40},{ty} {tx+tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs+42}" y="{ty-8}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'up_loop':
            xs = fx - fw//2
            pts = f'{xs},{fy} {xs-50},{fy} {xs-50},{ty} {tx},{ty+th//2}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
        elif direction == 'left_estop':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#DC2626" stroke-width="2" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-8}" text-anchor="middle" font-size="10" fill="#DC2626" font-weight="bold">{label}</text>')

    lines.append('</svg>')
    return '\n'.join(lines)


def make_docx(name, prefix, io_table):
    """Generate .docx report using python-docx"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = '宋体'

    # ---- Title Page ----
    for _ in range(4):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('沈阳城市建设学院')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('实训课程任务书')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = '黑体'

    doc.add_paragraph('')

    info_items = [
        ('实训课程名称', '电气控制实训'),
        ('实训课程题目', name),
        ('学期', '2026年春季学期'),
        ('学    院', '信息与控制工程学院'),
        ('专    业', '自动化（专升本）'),
        ('班    级', '25级 2班'),
        ('指导教师', '王丹阳    职称   讲师'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.name = '宋体'

    doc.add_page_break()

    # ---- Section 1: 设计任务与要求 ----
    h = doc.add_heading('一、设计任务与要求', level=1)
    for run in h.runs: run.font.name = '黑体'

    tasks = {
        '混凝土搅拌控制系统': [
            '设计一款基于可编程控制器（PLC）的混凝土搅拌控制系统，运用PLC、变频器、按钮指示灯等设备',
            '实现电机的启动、停止、正转、反转功能，速度可调（15Hz、20Hz、30Hz三档）',
            '模拟混凝土搅拌控制系统的搅拌运行过程：正转6秒→反转6秒→反复循环',
            '系统考虑安全信号，在异常情况下可使用急停按钮停止系统运行',
            '按下停止键后，电机延时3秒停止转动',
            '通过触摸屏显示运行/停止状态，并可控制系统的启动停止',
            '具备电气画图能力和PLC编程设计能力',
        ],
        '简易轮毂打紧机控制系统': [
            '设计一款基于可编程控制器（PLC）的简易轮毂打紧机控制系统，运用PLC、变频器、按钮指示灯等设备',
            '实现电机的启动、停止、正转、反转功能，速度可调',
            '按正转键，电机正转，正转指示灯点亮，到达指定位置停止',
            '按反转键，电机反转，反转指示灯点亮，到达指定位置停止',
            '4个流水灯循环点亮，1秒周期',
            '按下停止键，电机停止转动，灯熄灭',
            '系统考虑安全信号，异常情况急停按钮停止系统',
            '通过触摸屏显示运行/停止状态',
        ],
        '电动葫芦控制系统': [
            '设计一款基于可编程控制器（PLC）的电动葫芦控制系统，运用PLC、变频器、按钮指示灯等设备',
            '实现电机的启动、停止、正转（上升）、反转（下降）功能，速度可调',
            '按上升键，电机正转（上升），上升指示灯点亮，到达指定位置停止',
            '按下降键，电机反转（下降），下降指示灯点亮，到达指定位置停止',
            '6个流水灯循环点亮，1秒周期',
            '按下停止键，电机停止转动，灯熄灭',
            '系统考虑安全信号，异常情况急停按钮停止系统',
            '通过触摸屏显示运行/停止状态',
        ],
    }

    for t in tasks.get(name, tasks['混凝土搅拌控制系统']):
        p = doc.add_paragraph(t, style='List Bullet')

    # ---- Section 2: 器件选型 ----
    doc.add_heading('二、器件选型', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    components = [
        ['可编程控制器PLC', '三菱 FX3U-32MR/ES-A', '1台', '16入/16出，继电器输出，AC电源'],
        ['变频器', '三菱 FR-D720S-0.75K-CHT', '1台', '0.75kW，单相220V输入，多段速控制'],
        ['三相异步电机', 'YS-7124', '1台', '0.37kW，380V，1400rpm'],
        ['断路器', 'DZ47-63 3P C10', '1个', '主电路短路保护'],
        ['熔断器', 'RT18-32 6A', '2个', '控制电路过载保护'],
        ['交流接触器', 'CJX2-0910 220V', '2个', '正反转控制，含机械互锁'],
        ['热继电器', 'JR36-20 1-1.6A', '1个', '电机过载保护'],
        ['启动按钮', 'LAY39-11BN 绿色', '1个', '自复位常开触点'],
        ['停止按钮', 'LAY39-11BN 红色', '1个', '自复位常闭触点'],
        ['急停按钮', 'LAY39-11ZS 红色', '1个', '旋转复位，常闭触点'],
        ['速度选择按钮', 'LAY39-11BN 黄色', '3个', '15Hz/20Hz/30Hz选择'],
        ['运行指示灯', 'AD16-22DS 绿色', '1个', '运行状态指示'],
        ['正转指示灯', 'AD16-22DS 绿色', '1个', '正转状态指示'],
        ['反转指示灯', 'AD16-22DS 黄色', '1个', '反转状态指示'],
        ['触摸屏', '威纶通 TK6071iQ', '1台', '7寸，RS485通信，显示状态+控制'],
        ['开关电源', 'S-100-24', '1个', 'DC24V 4.5A，PLC及触摸屏供电'],
        ['导线及辅材', 'BVR 1.5mm² 等', '若干', '控制线及动力线'],
    ]

    table = doc.add_table(rows=len(components)+1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['序号', '器件名称', '型号规格', '数量', '备注']):
        cell = table.rows[0].cells[j]
        cell.text = hdr
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for i, row_data in enumerate(components):
        table.rows[i+1].cells[0].text = str(i+1)
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j+1].text = val
            for p in table.rows[i+1].cells[j+1].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # ---- Section 3: I/O地址分配 ----
    doc.add_heading('三、I/O地址分配', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    io_header = ['输入端', '说明', '输出端', '说明', '备注']
    io_table_obj = doc.add_table(rows=len(io_table)+1, cols=5)
    io_table_obj.style = 'Table Grid'
    io_table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, hdr in enumerate(io_header):
        cell = io_table_obj.rows[0].cells[j]
        cell.text = hdr
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for i, row_data in enumerate(io_table):
        for j, val in enumerate(row_data):
            io_table_obj.rows[i+1].cells[j].text = val
            for p in io_table_obj.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    # ---- Section 4: 电路分析 ----
    doc.add_heading('四、电路分析', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    doc.add_heading('4.1 主电路分析', level=2)
    circuit_analysis = [
        '主电路由断路器QF引入三相380V电源，经熔断器FU1保护后，通过两个交流接触器KM1（正转）和KM2（反转）控制三相异步电机。KM1与KM2之间设置机械互锁和电气互锁，防止两个接触器同时吸合导致相间短路。热继电器FR串接在主电路中，提供电机过载保护。变频器VFD输入端经接触器连接工频电源，输出端连接电机，通过多段速端子（S1-S3）选择运行频率。',
        '当KM1吸合时，电机正转（搅拌正方向）；当KM2吸合时，电源相序改变，电机反转。变频器通过预先设定的三段频率参数（15Hz/20Hz/30Hz），由PLC的数字量输出点控制多段速端子组合，实现速度切换。',
    ]
    for t in circuit_analysis:
        doc.add_paragraph(t)

    doc.add_heading('4.2 控制电路分析', level=2)
    control_analysis = [
        '控制电路采用DC24V供电，由开关电源将AC220V转换为DC24V。PLC的输入端口X000~X005接收按钮信号（启动、停止、急停、三个速度选择按钮）。PLC的输出端口Y000~Y004和Y010~Y012分别控制正反转接触器线圈、指示灯及变频器多段速端子。控制电路与主电路通过接触器实现电气隔离。',
        'PLC程序采用循环扫描方式工作。在每个扫描周期中，PLC先读取所有输入端状态，然后按梯形图顺序从上到下、从左到右执行逻辑运算，最后刷新所有输出端状态。正反转之间通过输出继电器Y000和Y001的常闭触点实现软件互锁（梯级4、5），配合外部接触器的机械互锁，形成双重保护，确保绝对安全。',
    ]
    for t in control_analysis:
        doc.add_paragraph(t)

    # ---- Section 5: 梯形图程序设计 ----
    doc.add_heading('五、梯形图程序设计', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    ladder_explanations = {
        '混凝土搅拌控制系统': [
            ('梯级1：运行使能（自锁电路）', '按下X000（启动）后，辅助继电器M0得电并自锁。X001（停止，常闭触点）和X002（急停，常闭触点）接通时自锁维持。当停止或急停按钮断开时，M0失电，系统停止运行。'),
            ('梯级2-3：正反转6秒循环', 'M0运行时，T0开始计时（K60=6秒）。T0未到期间，M1（正转中）得电→正转运行。T0计时到后，常闭T0断开M1（正转结束），同时T1开始计时（K60=6秒）→反转运行。T1计时到后，常闭T1断开，T0重新开始计时，如此循环往复，实现正转6秒→反转6秒→反复运行。'),
            ('梯级4-5：正反转输出（含互锁）', '梯级4：M0运行且T0未到（正转中）且Y001互锁未接通→Y000正转输出。梯级5：T0已到且T1未到（反转中）且Y000互锁未接通→Y001反转输出。T2（停止延时中）断开正反转回路，实现延时停止。正反转通过Y000/Y001的常闭触点实现电气互锁。'),
            ('梯级6：停止延时3秒', '按下X001（停止）→T2开始延时K30（3秒）。3秒后T2常开触点闭合→M10置位→梯级4/5中T2常闭触点断开→正反转输出停止。实现"按下停止→延时3秒→电机停止"的要求。'),
            ('梯级7-8：三段速度选择', 'X003（15Hz）→M3自锁；X004（20Hz）→M4自锁（与M3、M5互锁）；X005（30Hz）→M5自锁（与M3、M4互锁）。梯级8将M3/M4/M5映射到Y010/Y011/Y012（变频器多段速端子S1/S2/S3），分别输出15Hz、20Hz、30Hz。'),
            ('梯级9：状态指示灯', 'Y000（正转）→Y003正转绿灯亮；Y001（反转）→Y004反转黄灯亮；M0（运行）→Y002运行绿灯亮。'),
        ],
        '简易轮毂打紧机控制系统': [
            ('梯级1-2：正转/反转运行', '按下X000（正转键），M0自锁，Y000正转输出，Y002正转指示灯亮。按下X001（反转键），M1自锁，Y001反转输出，Y003反转指示灯亮。正反转之间通过Y000/Y001常闭触点互锁。'),
            ('梯级3：位置检测停止', 'X004（正转到位限位）常闭触点断开→Y000正转停止。X005（反转到位限位）常闭触点断开→Y001反转停止。到达指定位置后自动停止，无需手动干预。'),
            ('梯级4：停止/急停', 'X002（停止）或X003（急停）→M0/M1复位→所有输出停止。急停具有最高优先级，直接切断所有输出回路。'),
            ('梯级5：4流水灯（1s周期）', 'M2（运行中，M0或M1接通）→T0/T1交替振荡（各0.5s，1s周期）→计数器C0对T0上升沿计数（0→3循环）→C0值解码输出Y4→Y5→Y6→Y7→Y4循环，形成4位流水灯效果。'),
            ('梯级6-7：速度调节', 'X006（加速）→INCP D0；X007（减速）→DECP D0。D0=1/2/3分别对应低速/中速/高速。梯级7将D0值比较后输出到Y010/Y011/Y012（变频器多段速）。'),
        ],
        '电动葫芦控制系统': [
            ('梯级1-2：上升/下降运行', '按下X000（上升键），M0自锁，Y000正转（上升）输出，Y002上升指示灯亮。按下X001（下降键），M1自锁，Y001反转（下降）输出，Y003下降指示灯亮。正反转之间互锁。上限位X004断开→停止上升；下限位X005断开→停止下降。'),
            ('梯级3：位置检测停止', 'X004（上限位）串联在上升回路中，常闭→到达上限位时断开，Y000停止。X005（下限位）串联在下降回路中，常闭→到达下限位时断开，Y001停止。'),
            ('梯级4：停止/急停', 'X002（停止）或X003（急停）→M0/M1复位→所有输出停止。'),
            ('梯级5：6流水灯（1s周期）', 'M2（运行中）→T0/T1交替振荡（1s周期）→计数器C0（0→5循环）→解码输出Y4→Y5→Y6→Y7→Y10→Y11→Y4循环，形成6位流水灯效果。'),
            ('梯级6-7：速度调节', 'X006（加速）→INCP D0（上限3）；X007（减速）→DECP D0（下限1）。D0值比较后输出到Y012/Y013/Y014（变频器多段速1/2/3），分别对应低/中/高速。'),
        ],
    }

    for title, explanation in ladder_explanations.get(name, ladder_explanations['混凝土搅拌控制系统']):
        doc.add_heading(title, level=2)
        doc.add_paragraph(explanation)

    # Reference to ladder diagram file
    doc.add_paragraph(f'完整梯形图详见附件：{prefix}梯形图.svg')

    # ---- Section 6: 软件流程图 ----
    doc.add_heading('六、软件流程图', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    flow_desc = '系统上电后，首先执行初始化（参数清零、设置默认速度）。进入主循环后，依次扫描急停状态、启动按钮、正反转切换、速度调节按钮和停止按钮。程序采用典型的PLC循环扫描工作方式，每个扫描周期包括：输入采样→程序执行→输出刷新三个阶段。'
    doc.add_paragraph(flow_desc)
    doc.add_paragraph(f'完整软件流程图详见附件：{prefix}软件流程图.svg')

    # ---- Section 7: 系统功能说明 ----
    doc.add_heading('七、系统功能说明', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    func_descriptions = {
        '混凝土搅拌控制系统': [
            '启动功能：按下启动按钮（X000），系统运行使能M0置位，电机开始正转（Y000），正转指示灯（Y003）亮，运行指示灯（Y002）亮。',
            '自动循环搅拌：正转6秒（T0=K60）→自动切换反转6秒（T1=K60）→再次正转6秒……如此自动循环，模拟混凝土搅拌过程，无需人工干预。',
            '速度调节：通过三个速度选择按钮（X003/X004/X005）可独立选择15Hz（慢搅）、20Hz（中搅）、30Hz（快搅）。速度切换使用互锁逻辑，同一时间只能选择一种速度。',
            '停止延时：按下停止按钮（X001）后，T2开始3秒延时（K30）。3秒后常开T2触点闭合，断开正反转输出回路，电机平稳停止，防止急停造成物料飞溅。',
            '急停保护：任何时候按下急停按钮（X002），M0立即失电，所有输出（Y000~Y012）全部OFF，电机立刻停止，确保人员和设备安全。',
            '触摸屏监控：通过威纶通触摸屏（RS485通信），可实时显示系统运行/停止状态、当前转向、当前速度等级，并可触控操作。',
        ],
        '简易轮毂打紧机控制系统': [
            '正转打紧功能：按下正转键（X000），电机正转（Y000），正转指示灯（Y002）点亮，变频器驱动电机按设定转速正向旋转，进行轮毂打紧操作。到达指定位置（正转限位X004）后自动停止。',
            '反转松开功能：按下反转键（X001），电机反转（Y001），反转指示灯（Y003）点亮。到达指定位置（反转限位X005）后自动停止。',
            '速度可调：通过加速（X006）/减速（X007）按钮，可调节变频器输出频率，改变电机转速，适应不同规格轮毂的打紧力矩要求。',
            '4流水灯指示：系统运行时（M2=ON），4个指示灯（Y4→Y5→Y6→Y7→Y4）以1秒周期循环点亮，直观显示系统处于运行状态。',
            '停止与急停：按下停止键（X002）→所有输出复位；急停（X003）→最高优先级直接切断输出。限位开关同时作为自动停止信号。',
        ],
        '电动葫芦控制系统': [
            '上升功能：按下上升键（X000），电机正转（Y000），上升指示灯（Y002）点亮，电动葫芦提升重物。到达上限位（X004断开）自动停止，防止冲顶。',
            '下降功能：按下下降键（X001），电机反转（Y001），下降指示灯（Y003）点亮，电动葫芦下放重物。到达下限位（X005断开）自动停止，防止钢丝绳松脱。',
            '速度可调：通过加速/减速按钮改变变频器输出频率，实现快慢速升降，满足不同工况需求。',
            '6流水灯指示：系统运行时6个流水灯以1秒周期循环点亮，提升视觉效果，便于远距离判断系统运行状态。',
            '安全保护：急停按钮（X003）最高优先级→所有输出OFF；上限位（X004）/下限位（X005）作为正反向的极限保护；正反转之间电气互锁防止短路。',
        ],
    }

    for fd in func_descriptions.get(name, func_descriptions['混凝土搅拌控制系统']):
        doc.add_paragraph(fd, style='List Bullet')

    # ---- Section 8: 心得体会 ----
    doc.add_heading('八、心得体会', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    reflection = [
        f'通过本次《{name}》实训课程设计，我对可编程控制器（PLC）在工业自动化领域的应用有了更深入的理解。以下是我在实训过程中的几点心得体会：',
        '一、对PLC控制系统的整体认识：PLC作为工业自动化的核心控制器，通过输入采样、程序执行、输出刷新三个阶段的循环扫描工作方式，实现对工业设备的精确控制。本次设计使我掌握了PLC的I/O地址分配、梯形图编程以及变频器多段速控制等关键技术。',
        '二、梯形图编程的实践收获：梯形图是PLC最常用的编程语言，采用与继电器控制电路类似的表达方式，具有直观易懂的特点。在本次设计中，我重点掌握了以下编程技巧：(1)自锁电路实现运行使能的保持；(2)定时器实现延时控制和循环控制；(3)正反转互锁电路确保系统安全；(4)计数器实现流水灯循环效果。',
        '三、变频器多段速控制的理解：变频器通过改变输出频率来控制电机转速，多段速控制是最常用的控制方式之一。PLC通过数字量输出控制变频器的多段速端子（S1/S2/S3），实现预设频率的切换。',
        '四、安全设计的认识：工业控制系统安全是第一位的。本次设计从硬件（急停按钮、限位开关、接触器机械互锁）和软件（梯形图互锁逻辑、急停优先级）两个层面构建了完整的安全保护体系。这让我深刻认识到——在工程设计中，安全永远是第一优先级，没有安全就没有生产。',
        '五、收获与展望：本次实训将理论知识与工程实践紧密结合，从系统需求分析、器件选型、电路设计到梯形图编程、系统调试，完整经历了一个PLC控制系统的开发流程。这不仅加深了我对专业知识的理解，也培养了工程实践能力和解决实际问题的能力。在未来的学习和工作中，我将继续深入研究PLC高级应用、运动控制、工业通信网络等方向。',
    ]
    for t in reflection:
        doc.add_paragraph(t)

    # ---- Section 9: 参考文献 ----
    doc.add_heading('九、参考文献', level=1)
    for run in doc.paragraphs[-1].runs: run.font.name = '黑体'

    refs = [
        '三菱电机. FX3U系列微型可编程控制器用户手册[硬件篇]. 三菱电机株式会社, 2018.',
        '三菱电机. FX系列微型可编程控制器编程手册[基本指令篇]. 三菱电机株式会社, 2019.',
        '三菱电机. FR-D720S系列变频器使用手册. 三菱电机株式会社, 2020.',
        '廖常初. FX系列PLC编程及应用(第4版). 机械工业出版社, 2021.',
        '王曙光. 电气控制与PLC应用技术. 清华大学出版社, 2020.',
        '陈建明. 电气控制与PLC应用(第3版). 电子工业出版社, 2019.',
        '威纶通科技. TK6071iQ触摸屏用户手册. 威纶通科技有限公司, 2021.',
        'GB/T 15969.3-2017 可编程序控制器 第3部分: 编程语言[S]. 中国标准出版社, 2017.',
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    docx_path = os.path.join(OUT, f'{prefix}.docx')
    doc.save(docx_path)
    return docx_path


# ============================================================
# 主程序
# ============================================================
if __name__ == '__main__':
    print(f'=== 生成 {SYSTEM_NAME} ===')

    # Generate ladder SVG
    ladder_svg = make_ladder_svg(SYSTEM_NAME, LADDER_RUNGS, IO_TABLE)
    ladder_path = os.path.join(OUT, f'{FILE_PREFIX}梯形图.svg')
    with open(ladder_path, 'w', encoding='utf-8') as f:
        f.write(ladder_svg)
    print(f'[OK] 梯形图: {ladder_path} ({len(ladder_svg)} bytes)')

    # Generate flow chart SVG
    flow_svg = make_flow_svg(SYSTEM_NAME, FLOW_NODES, FLOW_CONNECTIONS)
    flow_path = os.path.join(OUT, f'{FILE_PREFIX}软件流程图.svg')
    with open(flow_path, 'w', encoding='utf-8') as f:
        f.write(flow_svg)
    print(f'[OK] 流程图: {flow_path} ({len(flow_svg)} bytes)')

    # Generate Word report
    docx_path = make_docx(SYSTEM_NAME, FILE_PREFIX, IO_TABLE)
    print(f'[OK] Word报告: {docx_path}')

    # XML validation
    for fp, fn in [(ladder_path, '梯形图'), (flow_path, '流程图')]:
        try:
            ET.fromstring(open(fp, 'r', encoding='utf-8').read())
            print(f'[OK] {fn} XML语法通过')
        except ET.ParseError as e:
            print(f'[FAIL] {fn} XML: {e}')

    print('完成!')
