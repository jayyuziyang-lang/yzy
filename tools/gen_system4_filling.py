#!/usr/bin/env python3
"""简易加注机控制系统 - 梯形图+流程图+Word报告 一键生成"""

import os, xml.etree.ElementTree as ET

OUT = r'D:\Desktop\CAD1'
os.makedirs(OUT, exist_ok=True)

SYSTEM_NAME = '简易加注机控制系统'
FILE_PREFIX = '简易加注机控制系统'
LIGHT_COUNT = 6  # 6流水灯
LIGHT_PERIOD = 1.5  # 1.5秒周期 → K75=0.75s每个

IO_TABLE = [
    ['X000', '启动按钮 SB1',    'Y000', '电机正转(加注) KM1', 'NO触点, 绿色'],
    ['X001', '停止按钮 SB2',    'Y001', '电机反转(回退) KM2', 'NC触点, 红色'],
    ['X002', '急停按钮 SB3',    'Y002', '运行指示灯 HL1',     '绿色'],
    ['X003', '加注完成 SQ1',    'Y003', '加注指示灯 HL2',     '蓝色'],
    ['X004', '反转/回退 SB4',   'Y004', '流水灯1 HL4',       '1.5s循环'],
    ['X005', '加速按钮 SB5',    'Y005', '流水灯2 HL5',       '1.5s循环'],
    ['X006', '减速按钮 SB6',    'Y006', '流水灯3 HL6',       '1.5s循环'],
    ['M0',   '加注使能(内部)',   'Y007', '流水灯4 HL7',       '1.5s循环'],
    ['M1',   '回退使能(内部)',   'Y010', '流水灯5 HL8',       '1.5s循环'],
    ['M2',   '系统运行中(内部)',  'Y011', '流水灯6 HL9',       '1.5s循环'],
    ['C0',   '流水灯计数器',     'Y012', '低速 VFD-S1',       '变频器多段速1'],
    ['D0',   '速度等级 1-3',    'Y013', '中速 VFD-S2',       '变频器多段速2'],
    ['',     '',                'Y014', '高速 VFD-S3',       '变频器多段速3'],
    ['',     '',                'T0/T1', '闪烁定时器',         'K75=0.75s, 1.5s周期'],
]

LADDER_RUNGS = [
    {
        'label': '梯级1: 加注使能(自锁, 到位自动停, 回退互锁)',
        'elements': [
            ('NO', 80, 'X000\n启动'),
            ('NC', 170, 'X003\n加注到位'),
            ('NC', 260, 'X001\n停止'),
            ('NC', 350, 'X002\n急停'),
            ('NC', 440, 'M1\n回退中'),
            ('COIL', 600, 'M0\n加注使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级2: 回退使能(手动回退, 加注互锁)',
        'elements': [
            ('NO', 80, 'X004\n回退键'),
            ('NC', 170, 'X001\n停止'),
            ('NC', 260, 'X002\n急停'),
            ('NC', 350, 'M0\n加注中'),
            ('COIL', 600, 'M1\n回退使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级3: 正转加注输出+回退输出(含互锁)',
        'elements': [
            ('NO', 80, 'M0\n加注使能'),
            ('NC', 170, 'Y001\n回退互锁'),
            ('COIL', 300, 'Y000\n加注正转'),
            ('COIL', 430, 'Y003\n加注灯'),
            ('NO', 530, 'M1\n回退使能'),
            ('NC', 620, 'Y000\n加注互锁'),
            ('COIL', 750, 'Y001\n回退'),
        ]
    },
    {
        'label': '梯级4: 运行中标志(加注或回退→M2运行)',
        'elements': [
            ('NO', 80, 'M0\n加注使能'),
            ('COIL', 200, 'M2\n运行中'),
            ('NO', 300, 'M1\n回退使能'),
        ]
    },
    {
        'label': '梯级5: 1.5s周期闪烁(振荡器T0⇄T1, K75=0.75s)',
        'elements': [
            ('NO', 80, 'M2\n运行中'),
            ('NC', 160, 'T1'),
            ('COIL', 250, 'T0\nK75'),
            ('NO', 340, 'T0'),
            ('COIL', 430, 'T1\nK75'),
        ]
    },
    {
        'label': '梯级6: 流水灯计数器(C0=0→1→2→3→4→5→0循环)',
        'elements': [
            ('NO', 80, 'T0\n0.75s'),
            ('FUNC', 170, 'INC\nC0'),
            ('FUNC', 260, '=\nC0 K6'),
            ('FUNC', 350, 'RST\nC0'),
        ]
    },
    {
        'label': '梯级7: 6流水灯解码输出(Y4→Y5→Y6→Y7→Y10→Y11)',
        'elements': [
            ('FUNC', 80, '=\nC0 K0'),
            ('COIL', 180, 'Y004\n灯1'),
            ('FUNC', 270, '=\nC0 K1'),
            ('COIL', 370, 'Y005\n灯2'),
            ('FUNC', 460, '=\nC0 K2'),
            ('COIL', 560, 'Y006\n灯3'),
            ('FUNC', 650, '=\nC0 K3'),
            ('COIL', 750, 'Y007\n灯4'),
        ]
    },
    {
        'label': '梯级8: 流水灯5-6解码(Y010/Y011)',
        'elements': [
            ('FUNC', 80, '=\nC0 K4'),
            ('COIL', 200, 'Y010\n灯5'),
            ('FUNC', 310, '=\nC0 K5'),
            ('COIL', 430, 'Y011\n灯6'),
        ]
    },
    {
        'label': '梯级9: 速度调节(加速/减速, D0范围1-3)',
        'elements': [
            ('NO', 80, 'X005\n加速'),
            ('FUNC', 180, 'INCP\nD0'),
            ('NO', 290, 'X006\n减速'),
            ('FUNC', 390, 'DECP\nD0'),
        ]
    },
    {
        'label': '梯级10: 速度输出至变频器(Y012低/Y013中/Y014高)',
        'elements': [
            ('FUNC', 80, '=\nD0 K1'),
            ('COIL', 180, 'Y012\n低速'),
            ('FUNC', 280, '=\nD0 K2'),
            ('COIL', 380, 'Y013\n中速'),
            ('FUNC', 480, '=\nD0 K3'),
            ('COIL', 580, 'Y014\n高速'),
        ]
    },
    {
        'label': '梯级11: 运行指示灯Y002',
        'elements': [
            ('NO', 80, 'M2\n运行中'),
            ('COIL', 200, 'Y002\n运行灯'),
        ]
    },
]


FLOW_NODES = [
    ('start', '开始\n(系统上电初始化)', 325, 30, 140, 45, 'startend', '#1A56DB'),
    ('init', '参数初始化\nD0=1(默认低速)\nC0=0', 325, 100, 140, 50, 'box', '#E8F0FE'),
    ('chk_estop', '急停按钮\n是否释放?', 325, 180, 110, 50, 'diamond', '#FFF3CD'),
    ('wait_cmd', '等待操作指令', 325, 260, 120, 45, 'box', '#E8F0FE'),
    ('chk_start', '启动按钮\n按下?', 200, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_rev_btn', '回退按钮\n按下?', 450, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('fill_run', 'M0=1 加注使能\nY000=1 正转加注\nY003=1 加注灯亮', 200, 430, 160, 55, 'box', '#D4EDDA'),
    ('rev_run', 'M1=1 回退使能\nY001=1 反转回退', 450, 430, 150, 50, 'box', '#FCE4D6'),
    ('run_loop', '运行中 M2=1\nY002运行灯亮\n6流水灯循环(1.5s)', 325, 520, 175, 55, 'box', '#D4EDDA'),
    ('chk_pos', '加注到位\nX003断开?', 325, 610, 120, 50, 'diamond', '#FFF3CD'),
    ('auto_stop', '加注完成!\nM0=0 自动停止\nY000/Y003=OFF', 200, 700, 165, 55, 'box', '#D4EDDA'),
    ('chk_stop', '停止按钮\n按下?', 450, 620, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_speed', '加速/减速\n按下?', 500, 710, 100, 50, 'diamond', '#FFF3CD'),
    ('adj_speed', '更新D0\n(INCP/DECP)', 500, 800, 120, 45, 'box', '#FCE4D6'),
    ('stop_act', '所有输出复位\nM0/M1=0\n加注停止', 150, 620, 140, 50, 'box', '#F8D7DA'),
    ('end', '结束', 150, 780, 100, 35, 'startend', '#6C757D'),
]

FLOW_CONNECTIONS = [
    ('start', 'init', 'down', ''),
    ('init', 'chk_estop', 'down', ''),
    ('chk_estop', 'wait_cmd', 'right', '是'),
    ('wait_cmd', 'chk_start', 'left', ''),
    ('wait_cmd', 'chk_rev_btn', 'right', ''),
    ('chk_start', 'fill_run', 'down', '是'),
    ('chk_start', 'wait_cmd', 'left_loop', '否'),
    ('chk_rev_btn', 'rev_run', 'down', '是'),
    ('chk_rev_btn', 'wait_cmd', 'right_loop', '否'),
    ('fill_run', 'run_loop', 'right', ''),
    ('rev_run', 'run_loop', 'left', ''),
    ('run_loop', 'chk_pos', 'left', ''),
    ('run_loop', 'chk_stop', 'right', ''),
    ('run_loop', 'chk_speed', 'down_right', ''),
    ('chk_pos', 'auto_stop', 'down', '到位'),
    ('chk_pos', 'run_loop', 'right_loop', '继续加注'),
    ('auto_stop', 'wait_cmd', 'up_loop', ''),
    ('chk_stop', 'stop_act', 'down', '是'),
    ('chk_stop', 'run_loop', 'right_loop', '否'),
    ('stop_act', 'end', 'down', ''),
    ('chk_speed', 'adj_speed', 'down', '是'),
    ('chk_speed', 'run_loop', 'right_loop', '否'),
    ('adj_speed', 'run_loop', 'up_loop', ''),
    ('chk_estop', 'stop_act', 'left_estop', '急停'),
]


def make_ladder_svg(name, rungs, io_table):
    W = 850
    rung_h = 78
    left_rail, right_rail, coil_x = 45, 745, 590
    rung_start_y, contact_y_offset = 55, 35
    total_h = len(rungs) * rung_h + 60 + 300
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {total_h}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<style>text{font-size:10px;fill:#1a1a1a}.io-label{font-size:8px;fill:#333;text-anchor:middle}.rung-label{font-size:9px;fill:#666;text-anchor:middle}.contact-no{fill:none;stroke:#1a1a1a;stroke-width:1.5}.func-box{fill:#f5f5f5;stroke:#1a1a1a;stroke-width:1.5}.wire{stroke:#1a1a1a;stroke-width:1.5;fill:none}.rail{fill:#333}</style>',
             f'<rect width="{W}" height="{total_h}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — PLC梯形图</text>',
             f'<text x="{W/2}" y="44" text-anchor="middle" font-size="10" fill="#666">PLC: 三菱FX3U | 变频器多段速调速(控制罐装量) | 6流水灯1.5s周期 | 加注到位自动停</text>']

    for i, rung in enumerate(rungs):
        y = rung_start_y + i * rung_h
        ey = y + contact_y_offset
        lines.append(f'<rect x="0" y="{y-3}" width="{W}" height="{rung_h-4}" fill="none" stroke="#e8e8e8" stroke-width="0.5"/>')
        lines.append(f'<rect x="{left_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<rect x="{right_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<line x1="{left_rail+3}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>')
        lines.append(f'<text x="{W/2}" y="{y+rung_h-6}" text-anchor="middle" class="rung-label">{rung["label"]}</text>')

        for elem in rung['elements']:
            typ, ex, lbl = elem
            if typ == 'NO':
                lines += [f'<line x1="{ex}" y1="{ey-10}" x2="{ex}" y2="{ey-16}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'NC':
                lines += [f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex+5}" y1="{ey-10}" x2="{ex-5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey-16}" x2="{ex}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'COIL':
                cx = ex
                lines += [f'<ellipse cx="{cx}" cy="{ey}" rx="13" ry="10" fill="none" stroke="#1a1a1a" stroke-width="1.5"/>',
                          f'<path d="M{cx-7},{ey-5} Q{cx-9},{ey} {cx-7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<path d="M{cx+7},{ey-5} Q{cx+9},{ey} {cx+7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<text x="{cx}" y="{ey+28}" class="io-label">{lbl}</text>',
                          f'<line x1="{cx+13}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>']
            elif typ == 'FUNC':
                fw, fh = 48, 28
                lines += [f'<rect x="{ex-fw/2}" y="{ey-fh/2}" width="{fw}" height="{fh}" rx="3" class="func-box"/>',
                          f'<text x="{ex}" y="{ey+4}" text-anchor="middle" font-size="9">{lbl}</text>']

        if rung.get('self_hold'):
            bx = rung['branch_from']
            sx = rung['self_x']
            by2 = ey + 46
            lines += [f'<line x1="{bx}" y1="{ey}" x2="{bx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{left_rail+6}" y1="{by2}" x2="{sx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{sx}" y1="{by2-12}" x2="{sx}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx}" y1="{by2+4}" x2="{sx}" y2="{by2+12}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2-4}" x2="{sx+5}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2+4}" x2="{sx+5}" y2="{by2+4}" class="contact-no"/>',
                      f'<text x="{sx}" y="{by2+26}" class="io-label">自锁</text>',
                      f'<line x1="{sx+10}" y1="{by2}" x2="{bx}" y2="{by2}" class="wire"/>']

    # I/O Table
    ty = rung_start_y + len(rungs) * rung_h + 12
    cw = [68, 100, 68, 100, 148]
    cx_pos = [25]
    for i in range(1, len(cw)): cx_pos.append(cx_pos[-1] + cw[i-1])
    lines.append(f'<text x="25" y="{ty}" font-size="13" font-weight="bold" fill="#1A56DB">I/O地址分配表 (三菱FX3U)</text>')
    tsy = ty + 12
    for j, (hdr, cx) in enumerate(zip(['输入端', '说明', '输出端', '说明', '备注'], cx_pos)):
        lines.append(f'<rect x="{cx}" y="{tsy}" width="{cw[j]}" height="20" fill="#1A56DB" stroke="#1A56DB"/>')
        lines.append(f'<text x="{cx+cw[j]/2}" y="{tsy+14}" text-anchor="middle" font-size="9" font-weight="bold" fill="white">{hdr}</text>')
    for ri, row in enumerate(io_table):
        ry = tsy + 20 + ri * 17
        fc = '#f8f9fa' if ri % 2 == 0 else '#fff'
        for j, (val, cx) in enumerate(zip(row, cx_pos)):
            lines.append(f'<rect x="{cx}" y="{ry}" width="{cw[j]}" height="17" fill="{fc}" stroke="#ddd" stroke-width="0.5"/>')
            lines.append(f'<text x="{cx+cw[j]/2}" y="{ry+12}" text-anchor="middle" font-size="8" fill="#333">{val}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_flow_svg(name, nodes, connections):
    W, H = 650, 900
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<defs><marker id="ah" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#333"/></marker></defs>',
             f'<rect width="{W}" height="{H}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — 软件流程图</text>',
             f'<text x="{W/2}" y="42" text-anchor="middle" font-size="10" fill="#666">主程序循环扫描 | 变频器三速(控制罐装量) | 6流水灯1.5s | 加注到位自动停</text>']

    nc = {}
    for n in nodes:
        nid, txt, cx, cy, w, h, shape, color = n
        nc[nid] = (cx, cy, w, h, shape)
        x, y = cx - w//2, cy - h//2
        if shape == 'startend':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="20" ry="20" fill="{color}" stroke="{color}"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="white" font-weight="bold">{ln}</text>')
        elif shape == 'diamond':
            lines.append(f'<polygon points="{cx},{y} {x+w},{cy} {cx},{y+h} {x},{cy}" fill="{color}" stroke="#856404" stroke-width="1.5"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#856404">{ln}</text>')
        elif shape == 'box':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="4" fill="{color}" stroke="#1A56DB" stroke-width="1.2"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#1a1a1a">{ln}</text>')

    for conn in connections:
        fid, tid, direction, label = conn
        fx, fy, fw, fh, _ = nc[fid]
        tx, ty, tw, th, _ = nc[tid]
        if direction == 'down':
            y1, y2 = fy + fh//2, ty - th//2
            lines.append(f'<line x1="{fx}" y1="{y1}" x2="{tx}" y2="{y2}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{fx+16}" y="{(y1+y2)/2+4}" font-size="9" fill="#856404">{label}</text>')
        elif direction in ('right', 'left'):
            x1 = fx + fw//2 if direction == 'right' else fx - fw//2
            x2 = tx - tw//2 if direction == 'right' else tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'down_right':
            y1 = fy + fh//2
            pts = f'{fx},{y1} {fx},{ty} {tx-tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{fx+16}" y="{(y1+ty)/2+4}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left_loop':
            xs = fx - fw//2; pts = f'{xs},{fy} {xs-25},{fy} {xs-25},{ty} {tx+tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs-28}" y="{fy-6}" font-size="9" fill="#DC2626">{label}</text>')
        elif direction == 'right_loop':
            xs = fx + fw//2; pts = f'{xs},{fy} {xs+28},{fy} {xs+28},{ty} {tx-tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs+30}" y="{ty-8}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'up_loop':
            xs = fx - fw//2; pts = f'{xs},{fy} {xs-45},{fy} {xs-45},{ty} {tx},{ty+th//2}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
        elif direction == 'left_estop':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#DC2626" stroke-width="2" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-8}" text-anchor="middle" font-size="10" fill="#DC2626" font-weight="bold">{label}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_docx(name, prefix, io_table):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.2)
    doc.styles['Normal'].font.size = Pt(11); doc.styles['Normal'].font.name = '宋体'

    for _ in range(3): doc.add_paragraph('')
    for txt, sz in [('沈阳城市建设学院', 20), ('实训课程任务书', 18)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.bold = True; r.font.name = '黑体'
    doc.add_paragraph('')
    for label, value in [('实训课程名称', '电气控制实训'), ('实训课程题目', name), ('学    院', '信息与控制工程学院'),
                          ('专    业', '自动化（专升本）'), ('班    级', '25级 2班'), ('指导教师', '王丹阳    职称   讲师'),
                          ('学    期', '2026年春季学期')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'{label}：{value}').font.size = Pt(13)
    doc.add_page_break()

    doc.add_heading('一、设计任务与要求', level=1)
    for t in ['设计基于PLC的简易加注机控制系统，运用PLC、变频器、按钮指示灯等设备',
              '按下启动键→电机正转加注，运行灯+加注灯亮，到达指定位置(加注完成)自动停止',
              '速度可调(三档变频)，用于控制罐装量(转速越高→加注越快→单位时间罐装量越大)',
              '6个流水灯循环点亮，1.5秒周期',
              '支持手动回退功能，用于加注完成后回退或维护操作',
              '停止键→电机停止、灯熄灭、停止加注；急停按钮→最高优先级立即停止',
              '通过触摸屏显示运行/停止状态及当前速度等级',
              '具备电气画图和PLC编程设计能力']:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('二、器件选型', level=1)
    comps = [
        ['PLC', '三菱 FX3U-32MR/ES-A', '1台', '16入/16出，继电器输出'],
        ['变频器', '三菱 FR-D720S-0.75K-CHT', '1台', '0.75kW，多段速控制'],
        ['三相异步电机', 'YS-7124', '1台', '0.37kW，380V'],
        ['断路器', 'DZ47-63 3P C10', '1个', '主电路短路保护'],
        ['交流接触器', 'CJX2-0910 220V', '2个', '加注/回退，机械互锁'],
        ['热继电器', 'JR36-20 1-1.6A', '1个', '过载保护'],
        ['启动按钮', 'LAY39-11BN 绿色', '1个', '自复位常开'],
        ['停止按钮', 'LAY39-11BN 红色', '1个', '自复位常闭'],
        ['急停按钮', 'LAY39-11ZS 红色', '1个', '旋转复位常闭'],
        ['回退按钮', 'LAY39-11BN 黄色', '1个', '自复位常开'],
        ['限位开关', 'LX19-001', '1个', '加注完成检测，常闭'],
        ['加注指示灯', 'AD16-22DS 蓝色', '1个', '加注状态指示'],
        ['流水指示灯', 'AD16-22DS 彩色', '6个', '6色，1.5s循环'],
        ['触摸屏', '威纶通 TK6071iQ', '1台', '7寸，RS485通信'],
        ['开关电源', 'S-100-24', '1个', 'DC24V 4.5A'],
    ]
    t = doc.add_table(rows=len(comps)+1, cols=5); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['序号','器件名称','型号规格','数量','备注']):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(comps):
        t.rows[i+1].cells[0].text = str(i+1)
        for j, v in enumerate(row):
            t.rows[i+1].cells[j+1].text = v
            for p in t.rows[i+1].cells[j+1].paragraphs:
                for r in p.runs: r.font.size = Pt(8)

    doc.add_heading('三、I/O地址分配', level=1)
    t2 = doc.add_table(rows=len(io_table)+1, cols=5); t2.style = 'Table Grid'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['输入端','说明','输出端','说明','备注']):
        t2.rows[0].cells[j].text = h
        for p in t2.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(io_table):
        for j, v in enumerate(row):
            t2.rows[i+1].cells[j].text = v if v else ''
            for p in t2.rows[i+1].cells[j].paragraphs:
                for r in p.runs: r.font.size = Pt(8)

    doc.add_heading('四、电路分析', level=1)
    doc.add_heading('4.1 主电路分析', level=2)
    doc.add_paragraph('主电路由断路器QF引入三相380V电源，经熔断器保护后，通过KM1(加注正转)和KM2(回退反转)控制电机。两接触器间设机械互锁+电气互锁，严防短路。热继电器FR提供过载保护。变频器VFD通过多段速端子(S1/S2/S3)接收PLC信号，输出三档频率控制加注泵转速——转速越高，加注越快，通过改变转速实现罐装量的精确调节。限位开关SQ1(加注完成)安装在加注工位，到达预设位置后常闭触点断开。')
    doc.add_heading('4.2 控制电路分析', level=2)
    doc.add_paragraph('控制电路采用DC24V供电，由开关电源转换。PLC输入X000~X006接收启动键、停止键、急停键、加注完成限位、回退键和加速/减速信号。输出端控制接触器、指示灯和变频器多段速。6个流水灯(1.5s周期)通过T0/T1振荡器(K75=0.75s)→计数器C0(0-5循环)→6个比较指令解码输出。加注限位X003(常闭)串联在加注使能回路中，到达位置后断开→M0失电→Y000加注自动停止。回退功能(X004手动触发)用于加注完成后回退或设备维护，与加注方向互为互锁。')

    doc.add_heading('五、梯形图程序设计', level=1)
    for title, exp in [
        ('梯级1-2：加注使能+回退使能',
         '梯级1：X000(启动)常开+M0自锁→串联X003(加注完成限位,NC)+X001(停止,NC)+X002(急停,NC)+M1常闭(回退互锁)→M0(加注使能)。到达加注位置→X003断开→M0失电自动停止。梯级2：X004(回退键)手动触发→M1自锁→回退运行，与M0互锁。'),
        ('梯级3：加注/回退输出',
         'M0→Y000(加注正转)+Y003(加注蓝灯)；M1→Y001(回退反转)。Y000与Y001常闭触点互锁，双重保障。'),
        ('梯级4-8：运行标志+1.5s振荡+6流水灯',
         'M0或M1→M2=ON(运行标志)。M2→T0/T1交替振荡(各K75=0.75s→总周期1.5s)。T0上升沿→C0加1(C0=0→1→2→3→4→5→0循环)。C0值解码→Y004~Y011依次点亮，6灯流水效果。'),
        ('梯级9-10：速度调节与VFD输出',
         'X005→INCP D0(上限3)；X006→DECP D0(下限1)。D0=1/2/3→Y012/Y013/Y014分别输出低/中/高速。加注速度直接影响罐装量——高速→大流量→短时间完成→罐装量由时间×速度决定。'),
        ('梯级11：运行指示灯', 'M2(运行中)→Y002运行绿灯亮。')]:
        doc.add_heading(title, level=2); doc.add_paragraph(exp)

    doc.add_heading('六、软件流程图', level=1)
    doc.add_paragraph('系统上电初始化(D0=1默认低速)→检测急停→等待操作指令→按启动键→M0自锁→Y000正转加注+Y003加注灯+6流水灯1.5s循环→检测加注完成限位X003→到位断开→M0失电自动停止→返回等待(可再次启动或回退)。运行中可调速。停止/急停→所有输出复位。')
    doc.add_paragraph(f'完整梯形图及流程图详见附件：{prefix}梯形图.svg、{prefix}软件流程图.svg')

    doc.add_heading('七、系统功能说明', level=1)
    for f in [
        '启动加注：按下启动键→电机正转加注，加注蓝灯+运行绿灯亮，6流水灯1.5s循环。',
        '自动停止：到达加注完成限位→X003断开→M0失电→Y000/Y003自动停止。实现精确罐装量控制。',
        '速度调节：加速/减速按钮三档调节变频器频率，控制加注泵转速，从而调节罐装量(转速∝加注速度)。',
        '手动回退：加注完成后按回退键→M1自锁→Y001反转回退→松手后自锁保持→到位后手动停。用于设备复位或维护。',
        '6流水灯：运行中6灯1.5s周期交替点亮，比常规1s周期更舒缓，适合工业现场远距离辨识。',
        '安全保护：正反转互锁(软件+硬件)、加注完成自动停机、急停最高优先级。']:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('八、心得体会', level=1)
    for t in [
        '通过"简易加注机控制系统"实训，我对PLC在流体罐装控制领域的应用有了深入理解。',
        '一、加注机的核心控制逻辑是"启停+到位检测"的线性流程，与搅拌机的"循环往复"和轮毂打紧机的"双向到位"不同，体现了工业设备控制逻辑的多样性。',
        '二、速度调节直接关联罐装量是本系统的一个特色。变频器频率越高→加注泵转速越快→单位时间罐装量越大。通过三档频率配合固定加注时间，可实现三种不同的罐装规格，这是"速度控制"转化为"计量控制"的工程思维。',
        '三、6流水灯1.5s周期的实现：T0/T1各K75=0.75s(比常规K5=0.5s慢50%)，C0计数0→5循环，6个比较解码。技术原理与4流水灯完全相同，体现了"振荡→计数→解码"范式的扩展性——改K值和K6即可适配任意周期和灯数。',
        '四、回退功能作为辅助操作独立于主控制流程，采用手动触发+自锁+互锁的设计，既不影响加注主流程，又能在需要时灵活使用，体现了"主辅分离"的程序设计思想。',
        '五、本次实训让我完整经历了从需求分析到程序设计的PLC工程实践，综合运用了定时器、计数器、比较指令、自锁互锁等核心技术，为从事工业自动化相关工作积累了宝贵经验。']:
        doc.add_paragraph(t)

    doc.add_heading('九、参考文献', level=1)
    for ref in ['三菱电机.FX3U系列用户手册[硬件篇].2018.','三菱电机.FX系列编程手册[基本指令篇].2019.','三菱电机.FR-D720S变频器使用手册.2020.','廖常初.FX系列PLC编程及应用(第4版).机械工业出版社,2021.','王曙光.电气控制与PLC应用技术.清华大学出版社,2020.','陈建明.电气控制与PLC应用(第3版).电子工业出版社,2019.','GB/T 15969.3-2017 可编程序控制器 第3部分:编程语言[S].']:
        doc.add_paragraph(ref)

    docx_path = os.path.join(OUT, f'{prefix}.docx')
    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    print(f'=== 生成 {SYSTEM_NAME} ===')

    ladder_svg = make_ladder_svg(SYSTEM_NAME, LADDER_RUNGS, IO_TABLE)
    ladder_path = os.path.join(OUT, f'{FILE_PREFIX}梯形图.svg')
    with open(ladder_path, 'w', encoding='utf-8') as f: f.write(ladder_svg)
    print(f'[OK] 梯形图: ({len(ladder_svg)} bytes)')

    flow_svg = make_flow_svg(SYSTEM_NAME, FLOW_NODES, FLOW_CONNECTIONS)
    flow_path = os.path.join(OUT, f'{FILE_PREFIX}软件流程图.svg')
    with open(flow_path, 'w', encoding='utf-8') as f: f.write(flow_svg)
    print(f'[OK] 流程图: ({len(flow_svg)} bytes)')

    docx_path = make_docx(SYSTEM_NAME, FILE_PREFIX, IO_TABLE)
    print(f'[OK] Word报告: {os.path.basename(docx_path)}')

    for fp, fn in [(ladder_path, '梯形图'), (flow_path, '流程图')]:
        try:
            ET.fromstring(open(fp, 'r', encoding='utf-8').read())
            print(f'[OK] {fn} XML语法通过')
        except ET.ParseError as e: print(f'[FAIL] {fn}: {e}')

    print('完成!')
