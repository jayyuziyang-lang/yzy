#!/usr/bin/env python3
"""电动葫芦控制系统 - 梯形图+流程图+Word报告 一键生成"""

import os, xml.etree.ElementTree as ET

OUT = r'D:\Desktop\CAD1'
os.makedirs(OUT, exist_ok=True)

SYSTEM_NAME = '电动葫芦控制系统'
FILE_PREFIX = '电动葫芦控制系统'

IO_TABLE = [
    ['X000', '上升按钮 SB1',   'Y000', '电机正转(上升) KM1', 'NO触点, 绿色'],
    ['X001', '下降按钮 SB2',   'Y001', '电机反转(下降) KM2', 'NO触点, 绿色'],
    ['X002', '停止按钮 SB3',   'Y002', '上升指示灯 HL1',    '绿色'],
    ['X003', '急停按钮 SB4',   'Y003', '下降指示灯 HL2',    '黄色'],
    ['X004', '上限位 SQ1',     'Y004', '流水灯1 HL4',      '1s循环'],
    ['X005', '下限位 SQ2',     'Y005', '流水灯2 HL5',      '1s循环'],
    ['X006', '加速按钮 SB5',   'Y006', '流水灯3 HL6',      '1s循环'],
    ['X007', '减速按钮 SB6',   'Y007', '流水灯4 HL7',      '1s循环'],
    ['M0',   '上升使能(内部)',  'Y010', '流水灯5 HL8',      '1s循环'],
    ['M1',   '下降使能(内部)',  'Y011', '流水灯6 HL9',      '1s循环'],
    ['M2',   '系统运行中(内部)', 'Y012', '低速 VFD-S1',      '变频器多段速1'],
    ['C0',   '流水灯计数器',    'Y013', '中速 VFD-S2',      '变频器多段速2'],
    ['D0',   '速度等级 1-3',    'Y014', '高速 VFD-S3',      '变频器多段速3'],
    ['',     '',               'T0/T1', '闪烁定时器',       'K5=0.5s, 1s周期'],
]

LADDER_RUNGS = [
    {
        'label': '梯级1: 上升使能(自锁, 上限位自动停, 下降互锁)',
        'elements': [
            ('NO', 80, 'X000\n上升键'),
            ('NC', 170, 'X004\n上限位'),
            ('NC', 260, 'X002\n停止'),
            ('NC', 350, 'X003\n急停'),
            ('NC', 440, 'M1\n下降中'),
            ('COIL', 600, 'M0\n上升使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级2: 下降使能(自锁, 下限位自动停, 上升互锁)',
        'elements': [
            ('NO', 80, 'X001\n下降键'),
            ('NC', 170, 'X005\n下限位'),
            ('NC', 260, 'X002\n停止'),
            ('NC', 350, 'X003\n急停'),
            ('NC', 440, 'M0\n上升中'),
            ('COIL', 600, 'M1\n下降使能'),
        ],
        'self_hold': True, 'self_x': 130, 'branch_from': 220
    },
    {
        'label': '梯级3: 上升/下降输出+指示灯(含互锁)',
        'elements': [
            ('NO', 80, 'M0\n上升使能'),
            ('NC', 170, 'Y001\n下降互锁'),
            ('COIL', 300, 'Y000\n上升'),
            ('COIL', 420, 'Y002\n上升灯'),
            ('NO', 520, 'M1\n下降使能'),
            ('NC', 610, 'Y000\n上升互锁'),
            ('COIL', 730, 'Y001\n下降'),
            ('COIL', 820, 'Y003\n下降灯'),
        ]
    },
    {
        'label': '梯级4: 运行中标志(M0上升或M1下降→M2运行)',
        'elements': [
            ('NO', 80, 'M0\n上升使能'),
            ('COIL', 200, 'M2\n运行中'),
            ('NO', 300, 'M1\n下降使能'),
        ]
    },
    {
        'label': '梯级5: 1s周期闪烁(振荡器T0⇄T1)',
        'elements': [
            ('NO', 80, 'M2\n运行中'),
            ('NC', 160, 'T1'),
            ('COIL', 240, 'T0\nK5'),
            ('NO', 320, 'T0'),
            ('COIL', 400, 'T1\nK5'),
        ]
    },
    {
        'label': '梯级6: 流水灯计数器(C0=0→1→2→3→4→5→0循环)',
        'elements': [
            ('NO', 80, 'T0\n0.5s脉冲'),
            ('FUNC', 170, 'INC\nC0'),
            ('FUNC', 260, '=\nC0 K6'),
            ('FUNC', 350, 'RST\nC0'),
        ]
    },
    {
        'label': '梯级7: 6流水灯解码输出(Y4→Y5→Y6→Y7→Y10→Y11)',
        'elements': [
            ('FUNC', 80, '=\nC0 K0'),
            ('COIL', 180, 'Y004\n灯1'),
            ('FUNC', 270, '=\nC0 K1'),
            ('COIL', 370, 'Y005\n灯2'),
            ('FUNC', 460, '=\nC0 K2'),
            ('COIL', 560, 'Y006\n灯3'),
            ('FUNC', 650, '=\nC0 K3'),
            ('COIL', 750, 'Y007\n灯4'),
        ]
    },
    {
        'label': '梯级8: 流水灯5-6解码(Y010/Y011)',
        'elements': [
            ('FUNC', 80, '=\nC0 K4'),
            ('COIL', 200, 'Y010\n灯5'),
            ('FUNC', 310, '=\nC0 K5'),
            ('COIL', 430, 'Y011\n灯6'),
        ]
    },
    {
        'label': '梯级9: 速度调节(加速/减速, D0范围1-3)',
        'elements': [
            ('NO', 80, 'X006\n加速'),
            ('FUNC', 180, 'INCP\nD0'),
            ('NO', 290, 'X007\n减速'),
            ('FUNC', 390, 'DECP\nD0'),
        ]
    },
    {
        'label': '梯级10: 速度输出至变频器(Y012低/Y013中/Y014高)',
        'elements': [
            ('FUNC', 80, '=\nD0 K1'),
            ('COIL', 180, 'Y012\n低速'),
            ('FUNC', 280, '=\nD0 K2'),
            ('COIL', 380, 'Y013\n中速'),
            ('FUNC', 480, '=\nD0 K3'),
            ('COIL', 580, 'Y014\n高速'),
        ]
    },
    {
        'label': '梯级11: 停止保护(未运行时批量复位输出)',
        'elements': [
            ('NC', 80, 'M2\n未运行'),
            ('FUNC', 180, 'ZRST\nY000 Y014'),
        ]
    },
]

FLOW_NODES = [
    ('start', '开始\n(系统上电初始化)', 325, 30, 140, 45, 'startend', '#1A56DB'),
    ('init', '参数初始化\nD0=1(默认低速)\nC0=0', 325, 100, 140, 50, 'box', '#E8F0FE'),
    ('chk_estop', '急停按钮\n是否释放?', 325, 180, 110, 50, 'diamond', '#FFF3CD'),
    ('wait_cmd', '等待操作指令', 325, 260, 120, 45, 'box', '#E8F0FE'),
    ('chk_up', '上升键\n按下?', 200, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_down', '下降键\n按下?', 450, 340, 100, 50, 'diamond', '#FFF3CD'),
    ('up_run', 'M0=1 上升\nY000=1 电机正转\nY002=1 上升灯亮', 200, 430, 155, 55, 'box', '#D4EDDA'),
    ('down_run', 'M1=1 下降\nY001=1 电机反转\nY003=1 下降灯亮', 450, 430, 155, 55, 'box', '#D4EDDA'),
    ('run_loop', '运行中 M2=1\n6流水灯循环(1s)', 325, 520, 160, 50, 'box', '#D4EDDA'),
    ('chk_limit', '到达限位\n(上/下)?', 325, 600, 110, 50, 'diamond', '#FFF3CD'),
    ('auto_stop', '自动停止\nM0/M1=0\nY000/Y001=OFF', 325, 680, 150, 50, 'box', '#F8D7DA'),
    ('chk_stop', '停止按钮\n按下?', 150, 680, 100, 50, 'diamond', '#FFF3CD'),
    ('chk_speed', '加速/减速\n按下?', 500, 680, 100, 50, 'diamond', '#FFF3CD'),
    ('adj_speed', '更新D0\n(INCP/DECP)', 500, 770, 120, 45, 'box', '#FCE4D6'),
    ('stop_act', '所有输出复位\nM0/M1=0', 150, 770, 130, 45, 'box', '#F8D7DA'),
    ('end', '结束', 150, 850, 100, 35, 'startend', '#6C757D'),
]

FLOW_CONNECTIONS = [
    ('start', 'init', 'down', ''),
    ('init', 'chk_estop', 'down', ''),
    ('chk_estop', 'wait_cmd', 'right', '是'),
    ('wait_cmd', 'chk_up', 'down', ''),
    ('wait_cmd', 'chk_down', 'down', ''),
    ('chk_up', 'up_run', 'down', '是'),
    ('chk_up', 'wait_cmd', 'left_loop', '否'),
    ('chk_down', 'down_run', 'down', '是'),
    ('chk_down', 'wait_cmd', 'right_loop', '否'),
    ('up_run', 'run_loop', 'right', ''),
    ('down_run', 'run_loop', 'left', ''),
    ('run_loop', 'chk_limit', 'down', ''),
    ('run_loop', 'chk_stop', 'left', ''),
    ('run_loop', 'chk_speed', 'right', ''),
    ('chk_limit', 'auto_stop', 'down', '是'),
    ('chk_limit', 'run_loop', 'right_loop', '否'),
    ('auto_stop', 'wait_cmd', 'up_loop', ''),
    ('chk_stop', 'stop_act', 'down', '是'),
    ('chk_stop', 'run_loop', 'left_loop', '否'),
    ('stop_act', 'end', 'down', ''),
    ('chk_speed', 'adj_speed', 'down', '是'),
    ('chk_speed', 'run_loop', 'right_loop', '否'),
    ('adj_speed', 'run_loop', 'up_loop', ''),
    ('chk_estop', 'stop_act', 'left_estop', '急停'),
]


def make_ladder_svg(name, rungs, io_table):
    W = 850
    rung_h = 78
    left_rail, right_rail, coil_x = 45, 745, 590
    rung_start_y, contact_y_offset = 55, 35
    total_h = len(rungs) * rung_h + 60 + 290
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {total_h}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<style>text{font-size:10px;fill:#1a1a1a}.io-label{font-size:8px;fill:#333;text-anchor:middle}.rung-label{font-size:9px;fill:#666;text-anchor:middle}.contact-no{fill:none;stroke:#1a1a1a;stroke-width:1.5}.func-box{fill:#f5f5f5;stroke:#1a1a1a;stroke-width:1.5}.wire{stroke:#1a1a1a;stroke-width:1.5;fill:none}.rail{fill:#333}</style>',
             f'<rect width="{W}" height="{total_h}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — PLC梯形图</text>',
             f'<text x="{W/2}" y="44" text-anchor="middle" font-size="10" fill="#666">PLC: 三菱FX3U | 变频器多段速 | 6流水灯1s周期 | 上下限位保护</text>']

    for i, rung in enumerate(rungs):
        y = rung_start_y + i * rung_h
        ey = y + contact_y_offset
        lines.append(f'<rect x="0" y="{y-3}" width="{W}" height="{rung_h-4}" fill="none" stroke="#e8e8e8" stroke-width="0.5"/>')
        lines.append(f'<rect x="{left_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<rect x="{right_rail}" y="{y}" width="3" height="{rung_h-12}" class="rail"/>')
        lines.append(f'<line x1="{left_rail+3}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>')
        lines.append(f'<text x="{W/2}" y="{y+rung_h-6}" text-anchor="middle" class="rung-label">{rung["label"]}</text>')

        for elem in rung['elements']:
            typ, ex, lbl = elem
            if typ == 'NO':
                lines += [f'<line x1="{ex}" y1="{ey-10}" x2="{ex}" y2="{ey-16}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'NC':
                lines += [f'<line x1="{ex-5}" y1="{ey-10}" x2="{ex+5}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex-5}" y1="{ey+10}" x2="{ex+5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex+5}" y1="{ey-10}" x2="{ex-5}" y2="{ey+10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey-16}" x2="{ex}" y2="{ey-10}" class="contact-no"/>',
                          f'<line x1="{ex}" y1="{ey+10}" x2="{ex}" y2="{ey+16}" class="contact-no"/>',
                          f'<text x="{ex}" y="{ey+28}" class="io-label">{lbl}</text>']
            elif typ == 'COIL':
                cx = ex
                lines += [f'<ellipse cx="{cx}" cy="{ey}" rx="13" ry="10" fill="none" stroke="#1a1a1a" stroke-width="1.5"/>',
                          f'<path d="M{cx-7},{ey-5} Q{cx-9},{ey} {cx-7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<path d="M{cx+7},{ey-5} Q{cx+9},{ey} {cx+7},{ey+5}" fill="none" stroke="#1a1a1a" stroke-width="1"/>',
                          f'<text x="{cx}" y="{ey+28}" class="io-label">{lbl}</text>',
                          f'<line x1="{cx+13}" y1="{ey}" x2="{right_rail}" y2="{ey}" class="wire"/>']
            elif typ == 'FUNC':
                fw, fh = 48, 28
                lines += [f'<rect x="{ex-fw/2}" y="{ey-fh/2}" width="{fw}" height="{fh}" rx="3" class="func-box"/>',
                          f'<text x="{ex}" y="{ey+4}" text-anchor="middle" font-size="9">{lbl}</text>']

        if rung.get('self_hold'):
            bx = rung['branch_from']
            sx = rung['self_x']
            by2 = ey + 46
            lines += [f'<line x1="{bx}" y1="{ey}" x2="{bx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{left_rail+6}" y1="{by2}" x2="{sx}" y2="{by2}" class="wire"/>',
                      f'<line x1="{sx}" y1="{by2-12}" x2="{sx}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx}" y1="{by2+4}" x2="{sx}" y2="{by2+12}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2-4}" x2="{sx+5}" y2="{by2-4}" class="contact-no"/>',
                      f'<line x1="{sx-5}" y1="{by2+4}" x2="{sx+5}" y2="{by2+4}" class="contact-no"/>',
                      f'<text x="{sx}" y="{by2+26}" class="io-label">M0 自锁</text>',
                      f'<line x1="{sx+10}" y1="{by2}" x2="{bx}" y2="{by2}" class="wire"/>']

    # I/O Table
    ty = rung_start_y + len(rungs) * rung_h + 12
    cw = [68, 100, 68, 100, 145]
    cx_pos = [25]
    for i in range(1, len(cw)): cx_pos.append(cx_pos[-1] + cw[i-1])
    lines.append(f'<text x="25" y="{ty}" font-size="13" font-weight="bold" fill="#1A56DB">I/O地址分配表 (三菱FX3U)</text>')
    tsy = ty + 12
    for j, (hdr, cx) in enumerate(zip(['输入端', '说明', '输出端', '说明', '备注'], cx_pos)):
        lines.append(f'<rect x="{cx}" y="{tsy}" width="{cw[j]}" height="20" fill="#1A56DB" stroke="#1A56DB"/>')
        lines.append(f'<text x="{cx+cw[j]/2}" y="{tsy+14}" text-anchor="middle" font-size="9" font-weight="bold" fill="white">{hdr}</text>')
    for ri, row in enumerate(io_table):
        ry = tsy + 20 + ri * 17
        fc = '#f8f9fa' if ri % 2 == 0 else '#fff'
        for j, (val, cx) in enumerate(zip(row, cx_pos)):
            lines.append(f'<rect x="{cx}" y="{ry}" width="{cw[j]}" height="17" fill="{fc}" stroke="#ddd" stroke-width="0.5"/>')
            lines.append(f'<text x="{cx+cw[j]/2}" y="{ry+12}" text-anchor="middle" font-size="8" fill="#333">{val}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_flow_svg(name, nodes, connections):
    W, H = 650, 960
    lines = [f'<?xml version="1.0" encoding="UTF-8"?>',
             f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {W} {H}" width="{W}" font-family="Microsoft YaHei, SimHei, sans-serif">',
             '<defs><marker id="ah" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#333"/></marker></defs>',
             f'<rect width="{W}" height="{H}" fill="#FAFBFC"/>',
             f'<text x="{W/2}" y="26" text-anchor="middle" font-size="15" font-weight="bold" fill="#1A56DB">{name} — 软件流程图</text>',
             f'<text x="{W/2}" y="42" text-anchor="middle" font-size="10" fill="#666">主程序循环扫描 | 变频器三速 | 6流水灯1s周期 | 上下限位保护</text>']

    nc = {}
    for n in nodes:
        nid, txt, cx, cy, w, h, shape, color = n
        nc[nid] = (cx, cy, w, h, shape)
        x, y = cx - w//2, cy - h//2
        if shape == 'startend':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="20" ry="20" fill="{color}" stroke="{color}"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="white" font-weight="bold">{ln}</text>')
        elif shape == 'diamond':
            lines.append(f'<polygon points="{cx},{y} {x+w},{cy} {cx},{y+h} {x},{cy}" fill="{color}" stroke="#856404" stroke-width="1.5"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#856404">{ln}</text>')
        elif shape == 'box':
            lines.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="4" fill="{color}" stroke="#1A56DB" stroke-width="1.2"/>')
            for li, ln in enumerate(txt.split('\n')):
                lines.append(f'<text x="{cx}" y="{cy-(len(txt.split(chr(10)))-1)*6+li*12+4}" text-anchor="middle" font-size="9" fill="#1a1a1a">{ln}</text>')

    for conn in connections:
        fid, tid, direction, label = conn
        fx, fy, fw, fh, _ = nc[fid]
        tx, ty, tw, th, _ = nc[tid]
        if direction == 'down':
            y1, y2 = fy + fh//2, ty - th//2
            lines.append(f'<line x1="{fx}" y1="{y1}" x2="{tx}" y2="{y2}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{fx+16}" y="{(y1+y2)/2+4}" font-size="9" fill="#856404">{label}</text>')
        elif direction in ('right', 'left'):
            x1 = fx + fw//2 if direction == 'right' else fx - fw//2
            x2 = tx - tw//2 if direction == 'right' else tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-6}" text-anchor="middle" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'left_loop':
            xs = fx - fw//2; pts = f'{xs},{fy} {xs-25},{fy} {xs-25},{ty} {tx+tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs-28}" y="{fy-6}" font-size="9" fill="#DC2626">{label}</text>')
        elif direction == 'right_loop':
            xs = fx + fw//2; pts = f'{xs},{fy} {xs+28},{fy} {xs+28},{ty} {tx-tw//2},{ty}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{xs+30}" y="{ty-8}" font-size="9" fill="#856404">{label}</text>')
        elif direction == 'up_loop':
            xs = fx - fw//2; pts = f'{xs},{fy} {xs-40},{fy} {xs-40},{ty} {tx},{ty+th//2}'
            lines.append(f'<polyline points="{pts}" fill="none" stroke="#333" stroke-width="1.5" marker-end="url(#ah)"/>')
        elif direction == 'left_estop':
            x1, x2 = fx - fw//2, tx + tw//2
            lines.append(f'<line x1="{x1}" y1="{fy}" x2="{x2}" y2="{ty}" stroke="#DC2626" stroke-width="2" marker-end="url(#ah)"/>')
            if label: lines.append(f'<text x="{(x1+x2)/2}" y="{fy-8}" text-anchor="middle" font-size="10" fill="#DC2626" font-weight="bold">{label}</text>')
    lines.append('</svg>')
    return '\n'.join(lines)


def make_docx(name, prefix, io_table):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    doc.styles['Normal'].font.size = Pt(12); doc.styles['Normal'].font.name = '宋体'

    for _ in range(4): doc.add_paragraph('')
    for txt, sz, b in [('沈阳城市建设学院', 22, True), ('实训课程任务书', 20, True)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.font.bold = b; r.font.name = '黑体'
    doc.add_paragraph('')
    for label, value in [('实训课程名称', '电气控制实训'), ('实训课程题目', name), ('学期', '2026年春季学期'),
                          ('学    院', '信息与控制工程学院'), ('专    业', '自动化（专升本）'),
                          ('班    级', '25级 2班'), ('指导教师', '王丹阳    职称   讲师')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'{label}：{value}').font.size = Pt(14)
    doc.add_page_break()

    doc.add_heading('一、设计任务与要求', level=1)
    for t in ['设计一款基于PLC的电动葫芦控制系统，运用PLC、变频器、按钮指示灯等设备',
              '实现电机启动、停止、正转（上升）、反转（下降）功能，速度可调',
              '按上升键→电机正转（上升），上升指示灯亮，到达上限位自动停止',
              '按下降键→电机反转（下降），下降指示灯亮，到达下限位自动停止',
              '6个流水灯循环点亮，1秒周期',
              '按下停止键→电机停止，灯熄灭；急停按钮→立即停止',
              '通过触摸屏显示运行/停止状态', '具备电气画图能力和PLC编程设计能力']:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('二、器件选型', level=1)
    components = [
        ['可编程控制器PLC', '三菱 FX3U-32MR/ES-A', '1台', '16入/16出，继电器输出'],
        ['变频器', '三菱 FR-D720S-0.75K-CHT', '1台', '0.75kW，多段速控制'],
        ['三相异步电机(起升)', 'ZDY12-4', '1台', '0.4kW，锥形转子制动'],
        ['断路器', 'DZ47-63 3P C10', '1个', '主电路短路保护'],
        ['熔断器', 'RT18-32 6A', '2个', '控制电路过载保护'],
        ['交流接触器', 'CJX2-0910 220V', '2个', '上升/下降控制，机械互锁'],
        ['热继电器', 'JR36-20 1-1.6A', '1个', '电机过载保护'],
        ['上升按钮', 'LAY39-11BN 绿色', '1个', '自复位常开触点'],
        ['下降按钮', 'LAY39-11BN 黄色', '1个', '自复位常开触点'],
        ['停止按钮', 'LAY39-11BN 红色', '1个', '自复位常闭触点'],
        ['急停按钮', 'LAY39-11ZS 红色', '1个', '旋转复位，常闭触点'],
        ['上限位开关', 'LX19-001', '1个', '防止冲顶，常闭触点'],
        ['下限位开关', 'LX19-001', '1个', '防止钢丝绳松脱，常闭触点'],
        ['流水指示灯', 'AD16-22DS 彩色', '6个', '红黄绿蓝白橙各1，1s循环'],
        ['上升指示灯', 'AD16-22DS 绿色', '1个', '上升状态指示'],
        ['下降指示灯', 'AD16-22DS 黄色', '1个', '下降状态指示'],
        ['触摸屏', '威纶通 TK6071iQ', '1台', '7寸，RS485通信'],
        ['开关电源', 'S-100-24', '1个', 'DC24V 4.5A'],
        ['导线及辅材', 'BVR 1.5mm² 等', '若干', '控制线及动力线'],
    ]
    table = doc.add_table(rows=len(components)+1, cols=5); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['序号', '器件名称', '型号规格', '数量', '备注']):
        table.rows[0].cells[j].text = hdr
        for p in table.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
    for i, row_data in enumerate(components):
        table.rows[i+1].cells[0].text = str(i+1)
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j+1].text = val
            for p in table.rows[i+1].cells[j+1].paragraphs:
                for run in p.runs: run.font.size = Pt(9)

    doc.add_heading('三、I/O地址分配', level=1)
    io_obj = doc.add_table(rows=len(io_table)+1, cols=5); io_obj.style = 'Table Grid'; io_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['输入端', '说明', '输出端', '说明', '备注']):
        io_obj.rows[0].cells[j].text = hdr
        for p in io_obj.rows[0].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.font.bold = True; run.font.size = Pt(9)
    for i, row_data in enumerate(io_table):
        for j, val in enumerate(row_data):
            io_obj.rows[i+1].cells[j].text = val
            for p in io_obj.rows[i+1].cells[j].paragraphs:
                for run in p.runs: run.font.size = Pt(8)

    doc.add_heading('四、电路分析', level=1)
    doc.add_heading('4.1 主电路分析', level=2)
    doc.add_paragraph('主电路由断路器QF引入三相380V电源，经熔断器FU1保护后，通过交流接触器KM1（上升/正转）和KM2（下降/反转）控制锥形转子制动电机。KM1与KM2间设置机械互锁和电气互锁，防止同时吸合导致相间短路。热继电器FR提供过载保护。变频器VFD通过多段速端子（S1-S3）输出不同频率，控制电动葫芦升降速度。锥形转子电机在断电时自动抱闸制动，防止重物下滑，满足起重设备的安全要求。')
    doc.add_heading('4.2 控制电路分析', level=2)
    doc.add_paragraph('控制电路采用DC24V供电。PLC输入端X000~X007接收上升键、下降键、停止键、急停键、上限位、下限位和加速/减速信号。输出端控制上升/下降接触器、指示灯和变频器。6个流水灯（Y004~Y007+Y010~Y011）通过计数器C0（0-5循环）和比较指令实现1s周期循环。上限位X004串联在上升回路防止冲顶，下限位X005串联在下降回路防止钢丝绳松脱，这是起重设备的关键安全措施。')

    doc.add_heading('五、梯形图程序设计', level=1)
    for title, exp in [
        ('梯级1-2：上升/下降使能（自锁+限位保护+互锁）',
         '按下X000（上升键）→M0自锁→Y000正转（上升）→上限位X004断开时自动停止。按下X001（下降键）→M1自锁→Y001反转（下降）→下限位X005断开时自动停止。M0与M1通过常闭触点互锁。X002（停止）和X003（急停）串联用于手动停止。'),
        ('梯级3：上升/下降输出+指示灯',
         'M0→Y000（上升）+Y002（上升绿灯）；M1→Y001（下降）+Y003（下降黄灯）。上升回路串联Y001常闭（下降互锁），下降回路串联Y000常闭（上升互锁），双重电气互锁保障安全。'),
        ('梯级4-6：运行标志+1s振荡+6流水灯计数',
         'M0或M1→M2=ON（运行标志）。M2→T0/T1交替振荡（K5=0.5s→1s周期）。T0上升沿→C0加1。C0=6时RST清零→C0在0→1→2→3→4→5→0之间循环。'),
        ('梯级7-8：6流水灯解码',
         'C0=0→Y004；C0=1→Y005；C0=2→Y006；C0=3→Y007；C0=4→Y010；C0=5→Y011。6个灯以1s间隔依次点亮形成流水效果。'),
        ('梯级9-10：速度调节与VFD输出',
         'X006→INCP D0（上限3）；X007→DECP D0（下限1）。D0=1→Y012低速；D0=2→Y013中速；D0=3→Y014高速。电动葫芦通常轻载高速、重载低速。'),
        ('梯级11：批量复位保护', 'M2未运行→ZRST Y000 Y014批量复位所有输出，确保停止后所有设备可靠关闭。'),
    ]: doc.add_heading(title, level=2); doc.add_paragraph(exp)
    doc.add_paragraph(f'完整梯形图详见附件：{prefix}梯形图.svg')

    doc.add_heading('六、软件流程图', level=1)
    doc.add_paragraph('系统上电初始化后进入主循环。检测急停→扫描上升/下降键→响应启动（含自锁+互锁+限位串接）→运行中6流水灯循环→检测限位（到达自动停）→检测停止键→检测速度调节。程序按PLC循环扫描方式执行。')
    doc.add_paragraph(f'完整软件流程图详见附件：{prefix}软件流程图.svg')

    doc.add_heading('七、系统功能说明', level=1)
    for fd in ['上升功能：按上升键（X000）→电机正转（Y000），上升绿灯（Y002）亮，电动葫芦提升重物。到达上限位（X004断开）自动停止，防止冲顶事故。',
               '下降功能：按下降键（X001）→电机反转（Y001），下降黄灯（Y003）亮，电动葫芦下放重物。到达下限位（X005断开）自动停止，防止钢丝绳松脱。',
               '速度可调：加速/减速按钮调节变频器三档频率（D0=1/2/3→Y012/Y013/Y014），适应不同重量物料的起升需求。',
               '6流水灯：运行中6个流水灯以1s周期循环点亮（Y4→Y5→Y6→Y7→Y10→Y11→Y4），远距离可辨。',
               '安全保护：上下限位（防冲顶/防松脱）、急停（最高优先级）、正反转互锁（软件+硬件）、锥形转子制动（断电自动抱闸）、未运行时批量复位输出。']:
        doc.add_paragraph(fd, style='List Bullet')

    doc.add_heading('八、心得体会', level=1)
    for t in [f'通过本次《{name}》实训课程设计，我对PLC在起重设备控制中的应用有了全面认识。',
              '一、电动葫芦的工作原理：电动葫芦是工厂车间最常用的起重设备，核心控制包括升降（正反转）、限位保护（上下限位）、制动保护（锥形转子电机断电抱闸）和速度调节（变频器）。本次设计完整覆盖了这些功能。',
              '二、6流水灯的实现：相比4流水灯，6流水灯需要更多的解码比较指令和输出端口。C0=0-5的6状态循环（INC+RST at K6）配合6个比较输出，实现6位流水效果，增强了对计数器+比较指令组合编程的理解。',
              '三、起重设备的安全设计：电动葫芦属于特种设备范畴，安全性要求极高。本次设计从限位开关（硬件）、互锁电路（硬件+软件）、急停（直接切断）、制动器（断电抱闸）、输出批量复位等多个层面构建了纵深防御体系，体现了"安全第一"的工程设计理念。',
              '四、ZRST批量复位指令的应用：梯级11使用ZRST（区域复位）指令批量复位Y000~Y014，相比逐个RST更高效，是三菱PLC编程的实用技巧。',
              '五、总结：本次实训让我掌握了PLC在起重设备中的典型控制方案——正反转+互锁+变频调速+位置检测+流水指示+多重安全保护。这些技术组合不仅适用于电动葫芦，也可推广到电梯、升降平台、堆垛机等各类垂直运动控制系统。']:
        doc.add_paragraph(t)

    doc.add_heading('九、参考文献', level=1)
    for ref in ['三菱电机. FX3U系列微型可编程控制器用户手册[硬件篇]. 2018.',
                 '三菱电机. FX系列微型可编程控制器编程手册[基本指令篇]. 2019.',
                 '三菱电机. FR-D720S系列变频器使用手册. 2020.',
                 '廖常初. FX系列PLC编程及应用(第4版). 机械工业出版社, 2021.',
                 '王曙光. 电气控制与PLC应用技术. 清华大学出版社, 2020.',
                 'GB 6067.1-2010 起重机械安全规程 第1部分: 总则[S]. 中国标准出版社, 2010.',
                 '威纶通科技. TK6071iQ触摸屏用户手册. 2021.',
                 'GB/T 15969.3-2017 可编程序控制器 第3部分: 编程语言[S]. 中国标准出版社, 2017.']:
        doc.add_paragraph(ref)

    docx_path = os.path.join(OUT, f'{prefix}.docx')
    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    print(f'=== 生成 {SYSTEM_NAME} ===')
    ladder_svg = make_ladder_svg(SYSTEM_NAME, LADDER_RUNGS, IO_TABLE)
    ladder_path = os.path.join(OUT, f'{FILE_PREFIX}梯形图.svg')
    with open(ladder_path, 'w', encoding='utf-8') as f: f.write(ladder_svg)
    print(f'[OK] 梯形图: {ladder_path} ({len(ladder_svg)} bytes)')

    flow_svg = make_flow_svg(SYSTEM_NAME, FLOW_NODES, FLOW_CONNECTIONS)
    flow_path = os.path.join(OUT, f'{FILE_PREFIX}软件流程图.svg')
    with open(flow_path, 'w', encoding='utf-8') as f: f.write(flow_svg)
    print(f'[OK] 流程图: {flow_path} ({len(flow_svg)} bytes)')

    docx_path = make_docx(SYSTEM_NAME, FILE_PREFIX, IO_TABLE)
    print(f'[OK] Word报告: {docx_path}')

    for fp, fn in [(ladder_path, '梯形图'), (flow_path, '流程图')]:
        try:
            ET.fromstring(open(fp, 'r', encoding='utf-8').read())
            print(f'[OK] {fn} XML语法通过')
        except ET.ParseError as e: print(f'[FAIL] {fn} XML: {e}')
    print('完成!')
