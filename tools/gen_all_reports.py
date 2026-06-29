#!/usr/bin/env python3
"""批量生成三个PLC控制系统的Word实训报告"""

import os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r'D:\Desktop\CAD1'
os.makedirs(OUT_DIR, exist_ok=True)

SCHOOL = '沈阳城市建设学院'
COLLEGE = '信息与控制工程学院'
MAJOR = '自动化（专升本）'
CLASS = '25级 2班'
TEACHER = '王丹阳'
TITLE = '讲师'
SEMESTER = '2026年 春季学期'
YEAR = '2026年5月'


def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_cover_page(doc, title_cn):
    """Add cover page"""
    # Blank lines for spacing
    for _ in range(3):
        doc.add_paragraph('')

    # 学号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('学  号')
    run.font.size = Pt(12)

    doc.add_paragraph('')

    # School name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SCHOOL)
    run.font.size = Pt(18)
    run.font.bold = True

    doc.add_paragraph('')

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('实训课程任务书')
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Course info
    info_items = [
        ('实训课程名称', '电气控制实训'),
        ('实训课程题目', title_cn),
        ('', f'（{SEMESTER}）'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label:
            run = p.add_run(f'{label}        ')
            run.font.size = Pt(14)
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.font.bold = True

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Academy info
    academy_items = [
        ('学    院', COLLEGE),
        ('专    业', MAJOR),
        ('班    级', CLASS),
        ('姓    名', ''),
        ('指导教师', f'{TEACHER}      职  称    {TITLE}'),
    ]
    for label, value in academy_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}         ')
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.size = Pt(12)

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'                     {YEAR}')
    run.font.size = Pt(12)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('沈阳城市建设学院教务处 制')
    run.font.size = Pt(11)

    doc.add_page_break()


def add_section_title(doc, title):
    """Add a section title"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_body_text(doc, text, bold=False, indent=False):
    """Add body text paragraph"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = hdr
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


# ============================================================
# 系统1: 混凝土搅拌控制系统 报告
# ============================================================
def gen_concrete_mixer_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title = '混凝土搅拌控制系统设计'
    add_cover_page(doc, title)

    # Task info table
    add_section_title(doc, '一、设计任务信息')

    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('设计时间', '2026年6月22日 — 2026年6月26日'),
        ('设计题目', '混凝土搅拌控制系统设计'),
    ]
    for label, value in info_data:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    # 任务内容
    add_section_title(doc, '二、任务内容')
    add_body_text(doc, '设计一款基于可编程控制器的混凝土搅拌控制系统。运用可编程控制器（PLC）、变频器、按钮、指示灯等电气元件，实现搅拌电机的启动、停止、正转、反转功能，并且搅拌速度可调，模拟混凝土搅拌控制系统的搅拌运行过程。同时系统需考虑安全信号，在异常情况下可使用急停按钮立即停止系统运行。要求学生具备电气绘图的能力、能够对可编程控制器进行编程设计的能力。', indent=True)

    # 设计条件
    add_section_title(doc, '三、设计条件')
    add_body_text(doc, '该设计要求学生设计一款速度可调的混凝土搅拌控制系统，能够通过触摸屏显示运行停止状态，并且控制系统的启动停止；该系统主要包括可编程控制器、变频器的设计部分。', indent=True)
    add_body_text(doc, '基本功能要求：', bold=True)
    add_body_text(doc, '（1）按下启动键后，状态指示灯显示相应状态，电机开始运行。', indent=True)
    add_body_text(doc, '（2）电机正转6秒，反转6秒，反复循环转动，模拟混凝土搅拌过程。', indent=True)
    add_body_text(doc, '（3）转速可以进行调节，分别设置为15Hz、20Hz和30Hz三档速度。', indent=True)
    add_body_text(doc, '（4）按下停止键后，电机延时3秒停止转动，保证搅拌桶内物料均匀排出。', indent=True)
    add_body_text(doc, '（5）系统配备急停按钮，紧急情况下可立即切断所有输出。', indent=True)

    # 成果要求
    add_section_title(doc, '四、成果要求')
    add_body_text(doc, '（1）完成电气接线并绘制电路原理图A3图纸1张。', indent=True)
    add_body_text(doc, '（2）设计PLC软件梯形图（见附件：梯形图_混凝土搅拌控制系统.svg）。', indent=True)
    add_body_text(doc, '（3）绘制软件流程图（见附件：软件流程图_混凝土搅拌控制系统.svg）。', indent=True)
    add_body_text(doc, '（4）编写实训报告3000字，包括系统方案、器件选型、功能说明、电路分析、心得体会等。', indent=True)

    # I/O分配
    add_section_title(doc, '五、I/O地址分配表')
    io_headers = ['输入端', '地址', '输出端', '地址', '说明']
    io_rows = [
        ['启动按钮 SB1', 'X000', '电机正转 KM1', 'Y000', '正转6秒'],
        ['停止按钮 SB2', 'X001', '电机反转 KM2', 'Y001', '反转6秒'],
        ['急停按钮 SB3', 'X002', '运行指示灯 HL1', 'Y002', '绿色'],
        ['速度选择1 SA1', 'X003', '正转指示灯 HL2', 'Y003', '绿色'],
        ['速度选择2 SA2', 'X004', '反转指示灯 HL3', 'Y004', '黄色'],
        ['-', '-', '15Hz输出 VFD-S1', 'Y010', '低速搅拌'],
        ['-', '-', '20Hz输出 VFD-S2', 'Y011', '中速搅拌'],
        ['-', '-', '30Hz输出 VFD-S3', 'Y012', '高速搅拌'],
    ]
    add_table(doc, io_headers, io_rows)

    doc.add_paragraph('')
    add_body_text(doc, '内部软元件：M0（运行使能）、M1（反转状态标志）、T0（正转定时器 K60=6s）、T1（反转定时器 K60=6s）、T2（停止延时定时器 K30=3s）。', indent=True)

    # 器件选型
    add_section_title(doc, '六、器件选型')
    add_body_text(doc, '（1）可编程控制器（PLC）：选用三菱FX3U-32MR/ES-A。该型号为继电器输出型，具有16点输入/16点输出，支持高速计数和脉冲输出，供电电源为AC 220V。FX3U系列是三菱电机的中型PLC，具有丰富的指令集，支持顺控程序、定时器、计数器、数据比较等功能，完全满足本系统控制需求。', indent=True)
    add_body_text(doc, '（2）变频器：选用三菱FR-D720S-0.75K-CHT。该变频器功率为0.75kW，输入电压为单相220V，输出频率0.2~400Hz，支持多段速运行（最多15段），具有过流、过压、欠压、过热等完善的保护功能。通过外部端子控制可实现多段速切换，本系统使用S1、S2、S3三个端子分别对应15Hz、20Hz、30Hz三档速度。', indent=True)
    add_body_text(doc, '（3）三相异步电动机：选用Y2-801-4型，功率0.55kW，额定电压380V，额定电流1.5A，额定转速1390r/min。采用Y型接法。', indent=True)
    add_body_text(doc, '（4）按钮：选用LAY39系列按钮开关。启动按钮为绿色（常开触点），停止按钮为红色（常闭触点），急停按钮为红色蘑菇头带自锁（常闭触点）。速度选择开关选用2位旋钮开关。', indent=True)
    add_body_text(doc, '（5）指示灯：选用AD16-22DS系列LED指示灯。运行指示灯绿色、正转指示灯绿色、反转指示灯黄色，工作电压为AC/DC 24V。', indent=True)
    add_body_text(doc, '（6）断路器：选用DZ47-63 C10型，额定电流10A，用于主回路短路和过载保护。', indent=True)
    add_body_text(doc, '（7）接触器：选用CJX2-0910型交流接触器，线圈电压AC 220V，额定电流9A，用于电机正反转主回路通断。', indent=True)
    add_body_text(doc, '（8）热继电器：选用NR2-25型，整定电流范围1.6~2.5A，用于电机过载保护。', indent=True)
    add_body_text(doc, '（9）开关电源：选用S-100-24型，输入AC 220V，输出DC 24V/4.5A，为PLC输入回路和指示灯供电。', indent=True)

    # 功能说明
    add_section_title(doc, '七、功能说明')
    add_body_text(doc, '本混凝土搅拌控制系统以PLC为核心控制器，实现对搅拌电机的正反转循环控制和速度调节。系统工作流程如下：', indent=True)
    add_body_text(doc, '（1）系统上电初始化：PLC上电后，所有输出复位，内部继电器M0、M1处于OFF状态，速度等级默认设置为D0=1（20Hz中速）。', indent=True)
    add_body_text(doc, '（2）启动运行：按下启动按钮（X000），M0置位并自锁，系统进入运行状态。M0=ON后，由于M1初始为OFF，Y000输出ON，电机正转接触器吸合，电机开始正转运行，同时T0开始计时6秒。运行指示灯Y002和正转指示灯Y003点亮。', indent=True)
    add_body_text(doc, '（3）正转→反转切换：T0计时满6秒后，置位M1（反转状态标志）。M1=ON后，断开正转回路（Y000=OFF），接通反转回路（Y001=ON），同时T1开始计时6秒。电机切换为反转运行，反转指示灯Y004点亮。', indent=True)
    add_body_text(doc, '（4）反转→正转切换：T1计时满6秒后，复位M1。M1=OFF后，断开反转回路（Y001=OFF），接通正转回路（Y000=ON），T0重新计时。如此反复循环，实现正转6秒→反转6秒→正转6秒→……的交替运行，模拟混凝土搅拌桶的双向搅拌过程。', indent=True)
    add_body_text(doc, '（5）速度调节：通过速度选择开关SA1（X003）和SA2（X004）的组合，可选择三档输出频率。X003=OFF且X004=OFF时Y010=ON（15Hz低速）；X003=ON且X004=OFF时Y011=ON（20Hz中速）；X004=ON时Y012=ON（30Hz高速）。变频器根据对应端子信号输出相应频率，实现搅拌速度的调节。', indent=True)
    add_body_text(doc, '（6）停止：按下停止按钮（X001，常闭触点断开），T2开始延时3秒计时。在T2计时的3秒内，电机继续运行（若处于正转则继续正转，若处于反转则继续反转）。T2计时到后，复位M0，所有输出断开，电机停止运行。延时停止的设计是为了让搅拌桶内物料能够均匀排出，避免突然停止造成物料残留。', indent=True)
    add_body_text(doc, '（7）急停保护：急停按钮（X002）串联在运行使能回路中，当急停按钮按下时，X002断开，M0立即失电，所有输出断开，系统紧急停止。急停按钮为红色蘑菇头带自锁型，复位时需旋转释放。', indent=True)
    add_body_text(doc, '（8）互锁保护：正转输出Y000和反转输出Y001之间存在电气互锁（梯形图中的常闭触点串联）和机械互锁（接触器的辅助常闭触点），确保正转和反转不会同时接通，防止主回路短路。', indent=True)

    # 电路分析
    add_section_title(doc, '八、电路分析')
    add_body_text(doc, '（1）主回路：三相电源L1、L2、L3经断路器QF1后，分别通过正转接触器KM1和反转接触器KM2接入电动机。KM1和KM2的主触点接线相序不同（其中两相对调），从而实现电机正反转。主回路中串联热继电器FR的发热元件，对电机进行过载保护。变频器VFD输入端经接触器KM3接入电源，输出端直接连接电机。', indent=True)
    add_body_text(doc, '（2）控制回路：PLC输出端Y000驱动中间继电器KA1（控制KM1正转接触器），Y001驱动中间继电器KA2（控制KM2反转接触器）。输出回路中串联急停按钮的常闭触点和热继电器的常闭触点，确保急停和过载时能可靠断开控制回路。', indent=True)
    add_body_text(doc, '（3）PLC输入回路：启动按钮（X000）、停止按钮（X001，常闭）、急停按钮（X002，常闭）、速度选择开关（X003、X004）均接入PLC输入端，采用DC 24V电源供电。停止按钮和急停按钮使用常闭触点的原因是：当按钮损坏或线路断线时，信号自然断开，系统自动停止，符合"故障-安全"原则。', indent=True)
    add_body_text(doc, '（4）变频器控制回路：PLC输出端Y010、Y011、Y012分别连接变频器的多功能输入端子S1、S2、S3。通过预先设置变频器参数（Pr.4=15Hz, Pr.5=20Hz, Pr.6=30Hz），当对应端子接通时，变频器输出预设频率。变频器正转信号由PLC输出Y000通过中间继电器控制。', indent=True)
    add_body_text(doc, '（5）电源回路：系统使用AC 220V为PLC、接触器线圈和开关电源供电。开关电源输出DC 24V为PLC输入回路和指示灯供电。变频器独立使用AC 220V电源。各部分电源独立配置断路器保护。', indent=True)
    add_body_text(doc, '（6）安全保护回路：系统设计了多重安全保护——急停按钮（硬件直接切断控制回路）、PLC内部逻辑互锁（正反转不能同时输出）、热继电器过载保护（电机过载时断开控制回路）、断路器短路保护（主回路和控制回路分别保护）。', indent=True)

    # 梯形图说明
    add_section_title(doc, '九、梯形图程序说明')
    add_body_text(doc, '梯形图共包含8个梯级（详见附件：梯形图_混凝土搅拌控制系统.svg），各梯级功能如下：', indent=True)
    add_body_text(doc, '梯级1：运行使能自锁电路。X000（启动）与M0自锁触点并联，串联X001（停止，常闭）、X002（急停，常闭）后驱动M0线圈。X000按下→M0=ON→自锁触点闭合→即使X000断开M0仍保持ON。停止或急停时断开回路→M0=OFF。', indent=True)
    add_body_text(doc, '梯级2：正转定时。M0=ON且M1=OFF时，T0开始计时K60（6秒）。', indent=True)
    add_body_text(doc, '梯级3：反转状态切换。T0计时到→SET M1（进入反转状态）→M1=ON→T1开始计时K60（6秒）→T1计时到→RST M1（回到正转状态），形成循环。', indent=True)
    add_body_text(doc, '梯级4：正转输出。M0=ON、M1=OFF（非反转状态）、Y001=OFF（反转互锁）三个条件同时满足时，Y000=ON（正转）。', indent=True)
    add_body_text(doc, '梯级5：反转输出。M0=ON、M1=ON（反转状态）、Y000=OFF（正转互锁）三个条件同时满足时，Y001=ON（反转）。', indent=True)
    add_body_text(doc, '梯级6：速度选择。X003和X004的常开/常闭组合判断，分别输出Y010（15Hz）、Y011（20Hz）、Y012（30Hz）。', indent=True)
    add_body_text(doc, '梯级7：停止延时。X001按下（常闭断开）→T2计时30×0.1s=3s→T2计时到→RST M0→系统停止。', indent=True)
    add_body_text(doc, '梯级8：状态指示灯。M0=ON→Y002（运行灯亮）；Y000=ON→Y003（正转灯亮）；Y001=ON→Y004（反转灯亮）。', indent=True)

    # 进度计划
    add_section_title(doc, '十、进度计划')
    progress_headers = ['时间', '设计内容']
    progress_rows = [
        ['6月22日', '介绍设计要求及任务，分析任务要求'],
        ['6月23日', '总体方案设计分析、绘制整体控制电路功能图，确定所需电气元件'],
        ['6月24日', '输入电路、输出电路的分析设计。软件流程图的设计及相关程序设计'],
        ['6月25日', 'PLC程序编写'],
        ['6月26日', '系统完善，编写实训报告并上交'],
    ]
    add_table(doc, progress_headers, progress_rows)

    # 心得体会
    add_section_title(doc, '十一、心得体会')
    add_body_text(doc, '通过本次混凝土搅拌控制系统的设计实训，我对PLC控制系统的设计有了全面而深入的理解。以下是我在实训过程中的几点体会：', indent=True)
    add_body_text(doc, '第一，PLC梯形图程序的设计需要清晰的逻辑思维。在设计搅拌控制程序时，我深刻体会到自锁电路是工业控制中最基础也最重要的环节——它不仅用于运行使能的保持，更关系到系统的安全启停。启动和停止按钮采用不同的触点类型（常开/常闭），是"故障-安全"原则的直接体现。', indent=True)
    add_body_text(doc, '第二，正反转控制中的互锁设计至关重要。在梯形图程序中，我在正转回路中串联了反转输出的常闭触点，在反转回路中也串联了正转输出的常闭触点，这样在任何情况下正转和反转都不会同时导通，从根本上避免了主回路相间短路的风险。这种"硬件+软件"双重互锁的设计思路值得在工业控制中广泛采用。', indent=True)
    add_body_text(doc, '第三，变频器的多段速控制让调速变得简单可靠。通过预先在变频器中设置各段频率参数（Pr.4~Pr.6），PLC只需控制三个数字输出端子即可实现三级调速，避免了模拟量控制的复杂性和干扰问题。15Hz适合低速精细搅拌，20Hz适合常规搅拌，30Hz适合快速搅拌，满足了不同工况需求。', indent=True)
    add_body_text(doc, '第四，延时停止的设计具有实际工程意义。混凝土搅拌完毕后如果立即停止，物料可能残留在搅拌桶内壁，而延时3秒让电机以当前速度继续运行一段时间，有助于物料的均匀排出。同时T2的延时设定值（K30）可以根据实际需要灵活调整，体现了PLC控制的灵活性。', indent=True)
    add_body_text(doc, '第五，本次设计让我认识到了工业控制系统与普通电气控制的区别。PLC控制不仅需要关注功能实现，更要关注安全保护——急停按钮的硬件切断、热继电器的过载保护、正反转的软硬件互锁、停止按钮的常闭接法——这些都是工业安全的基本要求。', indent=True)
    add_body_text(doc, '通过这次实训，我将课堂上学到的PLC编程理论、电气控制原理、变频器应用等知识综合运用到了实际设计中，不仅掌握了梯形图编程的基本方法和技巧，更理解了工业控制系统"安全第一"的设计理念。这对我今后从事自动化相关工作具有重要的指导意义。', indent=True)

    # Save
    path = os.path.join(OUT_DIR, '实训报告_混凝土搅拌控制系统.docx')
    doc.save(path)
    print(f'[OK] {path}')


# ============================================================
# 系统2: 简易轮毂打紧机控制系统 报告
# ============================================================
def gen_hub_tightener_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title = '简易轮毂打紧机控制系统设计'
    add_cover_page(doc, title)

    add_section_title(doc, '一、设计任务信息')
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    for label, value in [('设计时间', '2026年6月22日 — 2026年6月26日'), ('设计题目', title)]:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    add_section_title(doc, '二、任务内容')
    add_body_text(doc, '设计一款基于可编程控制器的简易轮毂打紧机控制系统。运用可编程控制器（PLC）、变频器、按钮、指示灯等电气元件，实现打紧电机的启动、停止、正转、反转功能，并且转速可调，模拟简易轮毂打紧机内的运行过程。同时系统需考虑安全信号，在异常情况下可使用急停按钮停止系统运行。要求学生具备电气绘图的能力、能够对可编程控制器进行编程设计的能力。', indent=True)

    add_section_title(doc, '三、设计条件')
    add_body_text(doc, '该设计要求学生设计一款速度可调的简易轮毂打紧机控制系统，能够通过触摸屏显示运行停止状态，并且控制系统的启动停止；该系统主要包括可编程控制器、变频器的设计部分。', indent=True)
    add_body_text(doc, '基本功能要求：', bold=True)
    add_body_text(doc, '（1）按下正转键，运行状态指示灯显示相应状态，电机正转，速度可调，正转指示灯点亮，到达指定位置（正转限位）时自动停止。', indent=True)
    add_body_text(doc, '（2）按下反转键，电机反转，反转指示灯点亮，到达指定位置（反转限位）时自动停止。', indent=True)
    add_body_text(doc, '（3）系统配备4个流水灯循环点亮，循环周期为1秒。', indent=True)
    add_body_text(doc, '（4）按下停止键，电机停止转动，所有指示灯熄灭。', indent=True)
    add_body_text(doc, '（5）系统配备急停按钮，紧急情况下可立即切断所有输出。', indent=True)

    add_section_title(doc, '四、成果要求')
    add_body_text(doc, '（1）完成电气接线并绘制电路原理图A3图纸1张。', indent=True)
    add_body_text(doc, '（2）设计PLC软件梯形图（见附件：梯形图_简易轮毂打紧机控制系统.svg）。', indent=True)
    add_body_text(doc, '（3）绘制软件流程图（见附件：软件流程图_简易轮毂打紧机控制系统.svg）。', indent=True)
    add_body_text(doc, '（4）编写实训报告3000字。', indent=True)

    add_section_title(doc, '五、I/O地址分配表')
    io_headers = ['输入端', '地址', '输出端', '地址', '说明']
    io_rows = [
        ['正转按钮 SB1', 'X000', '电机正转 KM1', 'Y000', '到达正转限位停止'],
        ['反转按钮 SB2', 'X001', '电机反转 KM2', 'Y001', '到达反转限位停止'],
        ['停止按钮 SB3', 'X002', '运行指示灯 HL1', 'Y002', '红色'],
        ['急停按钮 SB4', 'X003', '正转指示灯 HL2', 'Y003', '绿色'],
        ['正转限位 SQ1', 'X004', '反转指示灯 HL3', 'Y004', '黄色'],
        ['反转限位 SQ2', 'X005', '流水灯1 HL4', 'Y005', '1s周期循环'],
        ['加速按钮 SB5', 'X006', '流水灯2 HL5', 'Y006', '1s周期循环'],
        ['减速按钮 SB6', 'X007', '流水灯3 HL6', 'Y007', '1s周期循环'],
        ['-', '-', '流水灯4+低速', 'Y010', '灯4+速度1'],
        ['-', '-', '中速输出 VFD-S2', 'Y011', '中速'],
        ['-', '-', '高速输出 VFD-S3', 'Y012', '高速'],
    ]
    add_table(doc, io_headers, io_rows)
    doc.add_paragraph('')
    add_body_text(doc, '内部软元件：M1（正转使能）、M2（反转使能）、D0（速度等级0~2）、T10/T11（流水灯定时器 K2.5=0.25s，4灯×0.25s=1s周期）。', indent=True)

    add_section_title(doc, '六、器件选型')
    add_body_text(doc, '（1）PLC：三菱FX3U-32MR/ES-A，16入/16出，继电器输出型，AC 220V供电。', indent=True)
    add_body_text(doc, '（2）变频器：三菱FR-D720S-0.75K-CHT，0.75kW，单相220V输入，多段速控制。', indent=True)
    add_body_text(doc, '（3）电机：Y2-801-4型三相异步电动机，0.55kW，额定转速1390r/min。', indent=True)
    add_body_text(doc, '（4）限位开关：LX19-001型行程开关，一常开一常闭，用于正反转到位检测。当打紧机滑台到达正转极限位置时触发SQ1，到达反转极限位置时触发SQ2。', indent=True)
    add_body_text(doc, '（5）按钮：LAY39系列——正转按钮绿色（常开）、反转按钮黄色（常开）、停止按钮红色（常闭）、急停按钮红色蘑菇头（常闭）、加速/减速按钮黑色（常开）。', indent=True)
    add_body_text(doc, '（6）指示灯：AD16-22DS系列——运行灯红色、正转灯绿色、反转灯黄色、4个流水灯蓝色。', indent=True)
    add_body_text(doc, '（7）断路器：DZ47-63 C10型，10A，主回路保护。', indent=True)
    add_body_text(doc, '（8）接触器：CJX2-0910型，AC 220V线圈，9A，正反转控制。', indent=True)

    add_section_title(doc, '七、功能说明')
    add_body_text(doc, '本简易轮毂打紧机控制系统模拟轮毂螺栓的打紧和拆卸过程。正转对应螺栓打紧（拧紧），反转对应螺栓松动（拆卸）。系统核心功能如下：', indent=True)
    add_body_text(doc, '（1）正转打紧：按下正转按钮（X000），SET指令将M1置位（正转使能=ON）。在急停未按下（X003常闭）且反转未运行（Y001互锁触点闭合）的条件下，Y000=ON，电机正转启动，正转指示灯Y003点亮，4个流水灯（Y005/Y006/Y007/Y010）按1秒周期循环闪烁。当滑台到达正转极限位置时，正转限位开关SQ1（X004）触发，RST指令将M1复位→Y000=OFF→电机自动停止。', indent=True)
    add_body_text(doc, '（2）反转拆卸：按下反转按钮（X001），SET指令将M2置位（反转使能=ON）。在急停未按下且正转未运行的条件下，Y001=ON，电机反转启动，反转指示灯Y004点亮，流水灯继续运行。到达反转极限位置时，反转限位开关SQ2（X005）触发，RST M2→Y001=OFF→电机自动停止。', indent=True)
    add_body_text(doc, '（3）限位保护：正转限位SQ1（X004）和反转限位SQ2（X005）分别对应打紧机的两个极限位置。当任一限位被触发时，对应的使能继电器（M1或M2）被复位，电机自动停止。这种设计既保护了机械设备不超出行程范围，又实现了到达位置后自动停止的自动化控制。', indent=True)
    add_body_text(doc, '（4）4流水灯：当电机运行时（M1或M2=ON），T10/T11组成0.25秒脉冲发生器（T10=0.25s ON，T11=0.25s OFF）。四个流水灯Y005→Y006→Y007→Y010依次点亮，每个亮0.25秒，4个灯循环一周正好1秒，形成视觉上的"流水"效果。流水灯指示系统正在运行中。', indent=True)
    add_body_text(doc, '（5）速度调节：通过加速按钮（X006）和减速按钮（X007）调节速度等级D0（0~2）。INCP指令使D0递增（上限2），DECP指令使D0递减（下限0）。D0=1时Y010=ON（低速），D0=2时Y011=ON（中速），D0=0（或>=3）时Y012=ON（高速，默认）。速度可实时调节，无需停机。', indent=True)
    add_body_text(doc, '（6）停止与急停：按下停止按钮（X002）→同时复位M1和M2→电机停止、所有灯熄灭。按下急停按钮（X003）→串联在正反转输出回路中的X003常闭触点断开→Y000/Y001立即OFF→系统紧急停止。', indent=True)

    add_section_title(doc, '八、电路分析')
    add_body_text(doc, '（1）主回路：三相电源经断路器QF1后，通过正转接触器KM1或反转接触器KM2（两者互锁，不可同时闭合）接入电动机。KM1和KM2主触点采用不同相序接线实现正反转。热继电器FR串联在主回路中提供过载保护。变频器输入端经独立接触器控制，输出端连接电机。', indent=True)
    add_body_text(doc, '（2）控制回路：PLC输出Y000经中间继电器KA1控制KM1线圈；Y001经KA2控制KM2线圈。KM1和KM2的辅助常闭触点分别串联在对方线圈回路中，构成机械互锁。急停按钮的常闭触点和热继电器常闭触点串联在控制回路电源中。', indent=True)
    add_body_text(doc, '（3）限位开关回路：正转限位SQ1（X004）和反转限位SQ2（X005）接入PLC输入端。限位开关安装在打紧机滑台的两端极限位置。当滑台触碰到限位开关时，信号送入PLC，程序通过RST指令复位对应的使能继电器。限位开关采用常开触点（未触发时电路断开，触发时闭合），在梯形图中通过RST指令实现复位功能。', indent=True)
    add_body_text(doc, '（4）PLC输入回路：8个输入信号（正转、反转、停止、急停、正限位、反限位、加速、减速）均接入PLC输入端子X000~X007，采用DC 24V供电，输入电流约5~7mA。停止按钮和急停按钮使用常闭触点。', indent=True)
    add_body_text(doc, '（5）安全保护：硬件方面——急停按钮直接切断控制回路电源、热继电器过载保护、断路器短路保护、正反转接触器机械互锁；软件方面——梯形图中正反转输出电气互锁、限位自动停止保护。', indent=True)

    add_section_title(doc, '九、梯形图程序说明')
    add_body_text(doc, '梯形图共包含9个梯级（详见附件SVG文件），核心逻辑如下：', indent=True)
    add_body_text(doc, '梯级1-2：正反转使能控制。利用SET/RST指令实现——正转键SET M1，正转限位RST M1；反转键SET M2，反转限位RST M2。使用SET/RST而非自锁电路的原因是：正反转各有独立的停止条件（限位触发），用SET/RST更清晰。', indent=True)
    add_body_text(doc, '梯级3-4：正反转输出，带互锁和急停保护。M1条件驱动Y000正转，M2条件驱动Y001反转。', indent=True)
    add_body_text(doc, '梯级5：停止按钮复位M1和M2。', indent=True)
    add_body_text(doc, '梯级6：速度调节，INCP/DECP指令调整D0，比较指令控制速度输出。', indent=True)
    add_body_text(doc, '梯级7-8：4流水灯控制，T10/T11产生0.25s脉冲，交替点亮4个灯。', indent=True)
    add_body_text(doc, '梯级9：状态指示灯输出。', indent=True)

    add_section_title(doc, '十、进度计划')
    progress_headers = ['时间', '设计内容']
    progress_rows = [
        ['6月22日', '介绍设计要求及任务，分析任务要求'],
        ['6月23日', '总体方案设计分析、确定所需电气元件'],
        ['6月24日', '输入/输出电路分析设计，流程图及程序设计'],
        ['6月25日', 'PLC程序编写'],
        ['6月26日', '系统完善，编写实训报告'],
    ]
    add_table(doc, progress_headers, progress_rows)

    add_section_title(doc, '十一、心得体会')
    add_body_text(doc, '本次简易轮毂打紧机控制系统的设计实训，让我深入掌握了基于PLC的位置控制和顺序控制方法。以下是我在实训中的几点体会：', indent=True)
    add_body_text(doc, '第一，SET/RST指令在位置控制中的应用非常灵活。不同于自锁电路，SET/RST可以实现多条件启动和多条件停止（正转→正转键SET+限位RST+停止RST），使程序逻辑更加清晰。在轮毂打紧机这样需要精确到位停止的场景中，SET/RST配合限位开关是一个经典的解决方案。', indent=True)
    add_body_text(doc, '第二，流水灯的设计让我理解了定时器串联产生脉冲的原理。T10和T11组成一个自复位振荡器（T10=0.25s→T11=0.25s→循环），产生周期为0.5s的脉冲信号。4个流水灯通过在不同时刻响应这个脉冲，实现了依次亮灭的"流水"效果。这种技巧在工业设备的运行指示中非常常见。', indent=True)
    add_body_text(doc, '第三，限位开关在工业自动化中的重要性。限位开关不仅是位置检测元件，更是安全装置——当机械部件到达极限位置时，限位开关触发信号使PLC停止电机，防止机械碰撞和损坏。在设计时还需要考虑限位开关的安装位置和触发可靠性。', indent=True)
    add_body_text(doc, '第四，INCP/DECP脉冲执行指令的应用。使用P后缀（Pulse）确保每次按下按钮时D0只变化1个单位，而不是每个扫描周期都变化。如果没有P后缀，由于PLC扫描周期极短（几毫秒），一次按键可能导致D0直接加到上限或减到下限。', indent=True)
    add_body_text(doc, '第五，正反转互锁在轮毂打紧机中尤为重要。因为打紧机的正转和反转代表了完全相反的操作（拧紧螺栓vs松动螺栓），如果正反转同时接通不仅会导致电气短路，更会对机械结构造成严重冲击。通过梯形图互锁+接触器机械互锁+外部电路互锁三重保护，确保了系统的绝对安全。', indent=True)

    path = os.path.join(OUT_DIR, '实训报告_简易轮毂打紧机控制系统.docx')
    doc.save(path)
    print(f'[OK] {path}')


# ============================================================
# 系统3: 电动葫芦控制系统 报告
# ============================================================
def gen_electric_hoist_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    title = '电动葫芦控制系统设计'
    add_cover_page(doc, title)

    add_section_title(doc, '一、设计任务信息')
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    for label, value in [('设计时间', '2026年6月22日 — 2026年6月26日'), ('设计题目', title)]:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    add_section_title(doc, '二、任务内容')
    add_body_text(doc, '设计一款基于可编程控制器的电动葫芦控制系统。运用可编程控制器（PLC）、变频器、按钮、指示灯等电气元件，实现电动葫芦电机的启动、停止、正转（上升）、反转（下降），并且速度可调，模拟电动葫芦的运行过程。同时系统需考虑安全信号，在异常情况下可使用急停按钮停止系统运行。要求学生具备电气绘图的能力、能够对可编程控制器进行编程设计的能力。', indent=True)

    add_section_title(doc, '三、设计条件')
    add_body_text(doc, '该设计要求学生设计一款速度可调的电动葫芦控制系统，能够通过触摸屏显示运行停止状态，并且控制系统的启动停止；该系统主要包括可编程控制器、变频器的设计部分。', indent=True)
    add_body_text(doc, '基本功能要求：', bold=True)
    add_body_text(doc, '（1）按下上升键，运行状态指示灯显示相应状态，电机正转，上升指示灯点亮，到达指定位置（上限位）时自动停止。', indent=True)
    add_body_text(doc, '（2）按下下降键，电机反转，下降指示灯点亮，到达指定位置（下限位）时自动停止。', indent=True)
    add_body_text(doc, '（3）系统配备6个流水灯循环点亮，循环周期为1秒。', indent=True)
    add_body_text(doc, '（4）按下停止键，电机停止转动，所有指示灯熄灭。', indent=True)
    add_body_text(doc, '（5）系统配备急停按钮，紧急情况下可立即切断所有输出。', indent=True)

    add_section_title(doc, '四、成果要求')
    add_body_text(doc, '（1）完成电气接线并绘制电路原理图A3图纸1张。', indent=True)
    add_body_text(doc, '（2）设计PLC软件梯形图（见附件：梯形图_电动葫芦控制系统.svg）。', indent=True)
    add_body_text(doc, '（3）绘制软件流程图（见附件：软件流程图_电动葫芦控制系统.svg）。', indent=True)
    add_body_text(doc, '（4）编写实训报告3000字。', indent=True)

    add_section_title(doc, '五、I/O地址分配表')
    io_headers = ['输入端', '地址', '输出端', '地址', '说明']
    io_rows = [
        ['上升按钮 SB1', 'X000', '电机正转(上升)', 'Y000', '到达上限位停止'],
        ['下降按钮 SB2', 'X001', '电机反转(下降)', 'Y001', '到达下限位停止'],
        ['停止按钮 SB3', 'X002', '运行指示灯 HL1', 'Y002', '红色'],
        ['急停按钮 SB4', 'X003', '上升指示灯 HL2', 'Y003', '绿色'],
        ['上限位 SQ1', 'X004', '下降指示灯 HL3', 'Y004', '黄色'],
        ['下限位 SQ2', 'X005', '流水灯1 HL4', 'Y005', '1s周期循环'],
        ['加速按钮 SB5', 'X006', '流水灯2 HL5', 'Y006', '1s周期循环'],
        ['减速按钮 SB6', 'X007', '流水灯3 HL6', 'Y007', '1s周期循环'],
        ['-', '-', '流水灯4 HL7', 'Y010', '1s周期循环'],
        ['-', '-', '流水灯5 HL8', 'Y011', '1s周期循环'],
        ['-', '-', '流水灯6 HL9', 'Y012', '1s周期循环'],
    ]
    add_table(doc, io_headers, io_rows)
    doc.add_paragraph('')
    add_body_text(doc, '内部软元件：M1（上升使能）、M2（下降使能）、M0（运行状态，M1或M2=ON时ON）、D0（速度等级0~2）、T20/T21（流水灯定时器 K1.67=0.167s，6灯×0.167s≈1s周期）。', indent=True)

    add_section_title(doc, '六、器件选型')
    add_body_text(doc, '（1）PLC：三菱FX3U-32MR/ES-A，16入/16出，继电器输出型。', indent=True)
    add_body_text(doc, '（2）变频器：三菱FR-D720S-0.75K-CHT，0.75kW，多段速控制。电动葫芦的上升和下降需要不同的速度，上升时通常需要较慢速度以保证安全，下降时可适当加快。', indent=True)
    add_body_text(doc, '（3）电机：YEJ系列电磁制动三相异步电动机（锥形转子电机），功率0.75kW。电动葫芦专用电机带有电磁制动器——通电时制动器释放（电机可转动），断电时制动器抱闸（电机锁定）。这一特性确保了在停电或急停时，负载不会因重力而下坠，是电动葫芦安全运行的关键。', indent=True)
    add_body_text(doc, '（4）限位开关：LX19-001型行程开关。上限位SQ1安装于电动葫芦上升的极限高度位置，防止吊钩冲顶；下限位SQ2安装于下降的极限低位置，防止钢丝绳过度放出。这两个限位开关是电动葫芦最重要的安全装置。', indent=True)
    add_body_text(doc, '（5）按钮：LAY39系列——上升按钮绿色（常开）、下降按钮黄色（常开）、停止按钮红色（常闭）、急停按钮红色蘑菇头带自锁（常闭）。按钮布置采用"上绿下黄"的标准电动葫芦操作台布局。', indent=True)
    add_body_text(doc, '（6）指示灯：AD16-22DS系列LED指示灯。上升指示灯绿色、下降指示灯黄色、运行灯红色、6个流水灯蓝色，DC 24V供电。', indent=True)
    add_body_text(doc, '（7）断路器：DZ47-63 C10型，10A。', indent=True)
    add_body_text(doc, '（8）接触器：CJX2-0910型，9A，上升和下降各一个，带机械互锁。', indent=True)

    add_section_title(doc, '七、功能说明')
    add_body_text(doc, '本电动葫芦控制系统模拟工业生产中常见的起重提升设备。电动葫芦是一种轻小型起重设备，广泛用于工厂、仓库、码头的物料吊运。系统核心功能如下：', indent=True)
    add_body_text(doc, '（1）上升控制：按下上升按钮（X000）→SET M1（上升使能=ON）。在急停未按下且下降未运行（Y001互锁）的条件下，Y000=ON→电机正转（定义为上升方向）→吊钩上升。上升过程中Y003（上升指示灯，绿色）点亮，6个流水灯按1秒周期循环闪烁。当吊钩上升至上限位高度时，上限位开关SQ1（X004）被触发→RST M1→Y000=OFF→电机自动停止，防止吊钩冲顶。', indent=True)
    add_body_text(doc, '（2）下降控制：按下下降按钮（X001）→SET M2（下降使能=ON）。在急停未按下且上升未运行（Y000互锁）的条件下，Y001=ON→电机反转（定义为下降方向）→吊钩下降。下降过程中Y004（下降指示灯，黄色）点亮，流水灯运行。到达下限位时，下限位开关SQ2（X005）触发→RST M2→Y001=OFF→电机自动停止。', indent=True)
    add_body_text(doc, '（3）6流水灯设计：当电机运行时（M1或M2=ON→M0=ON），T20/T21组成的自复位振荡器产生周期为0.334秒的脉冲（T20=0.167s ON, T21=0.167s OFF）。6个流水灯Y005/Y006/Y007/Y010/Y011/Y012分别在T20的不同状态下点亮，每个灯亮约0.167秒后切换到下一个，6个灯循环一周约1秒。相比4流水灯，6流水灯的过渡更加平滑流畅，视觉上近似于连续的"跑马灯"效果，状态指示更加醒目。', indent=True)
    add_body_text(doc, '（4）速度调节：上升和下降可独立设定不同的速度。通过加速按钮（X006）和减速按钮（X007）调节D0（0~2级）。由于电动葫芦上升时安全要求更高，通常上升使用低速（Y011），下降可适当使用中速（Y012）。变频器的多段速功能使得这种灵活调速变得简单易行。', indent=True)
    add_body_text(doc, '（5）安全保护系统：电动葫芦的安全保护设计尤为严格——①上下限位开关防止吊钩超出行程范围（冲顶或钢丝绳过度放出）；②电磁制动器在断电时自动抱闸，防止负载自由坠落；③急停按钮可随时切断动力和控制电源；④PLC梯形图中的正反转互锁和限位保护。以上多重保护共同构成了电动葫芦的安全防护体系。', indent=True)
    add_body_text(doc, '（6）停止操作：按下停止按钮（X002）→RST M1和RST M2同时执行→电机停止、所有指示灯熄灭。此外，X003（急停）不仅通过RST指令复位M1/M2，还采用ZRST指令（区间复位）一次性复位M1~M2，并通过X003常闭触点从硬件层面切断正反转输出回路，实现软硬件双重急停保护。', indent=True)

    add_section_title(doc, '八、电路分析')
    add_body_text(doc, '（1）主回路：三相电源L1/L2/L3经断路器QF1后，分别通过上升接触器KM1和下降接触器KM2接入电动机。KM1和KM2采用不同的相序接线实现正反转。在电动机接线盒旁并联电磁制动器线圈YB——当KM1或KM2吸合时，YB同时得电，制动器释放；当两个接触器都断开时，YB失电，制动器在弹簧力作用下抱闸，锁住电机轴。', indent=True)
    add_body_text(doc, '（2）制动回路是关键安全环节。电磁制动器YB的线圈电压通常为AC 380V（取自电机端电压）或DC 24V（通过整流桥）。制动器的响应时间直接影响提升安全——通电释放和断电抱闸的时间差必须足够短（通常<0.1秒），以防止断电瞬间负载下坠。', indent=True)
    add_body_text(doc, '（3）控制回路：PLC输出Y000经KA1驱动KM1（上升），Y001经KA2驱动KM2（下降）。KM1和KM2的辅助常闭触点交叉串联构成机械互锁。急停按钮常闭触点串联在控制电源中。热继电器FR的常闭触点也串联其中，提供过载保护。', indent=True)
    add_body_text(doc, '（4）限位开关安装：上限位SQ1安装在电动葫芦轨道/钢丝绳卷筒的上端极限位置，当吊钩上升至最高安全高度时被触发；下限位SQ2安装在钢丝绳卷筒的下端极限位置（钢丝绳即将完全放出时触发）。限位开关的安装必须牢固可靠，且触发机构需要定期检查和维护。', indent=True)
    add_body_text(doc, '（5）电源回路：系统采用AC 220V为主控电源，经开关电源转换为DC 24V供PLC输入和指示灯。变频器独立使用AC 220V。由于电动葫芦可能在户外或潮湿环境中使用，电气柜防护等级建议不低于IP54，所有接线端子需防潮处理。', indent=True)
    add_body_text(doc, '（6）多重安全保护路径：第1层——限位开关（软件限制行程）；第2层——PLC程序互锁（正反转不可同时ON）；第3层——接触器机械互锁（KM1和KM2不可同时闭合）；第4层——急停按钮（硬件切断控制电源）；第5层——电磁制动器（断电自动抱闸）；第6层——断路器+热继电器（短路+过载保护）。六层保护使电动葫芦在各种异常情况下都能安全停止。', indent=True)

    add_section_title(doc, '九、梯形图程序说明')
    add_body_text(doc, '梯形图共包含9个梯级（详见附件SVG文件），核心逻辑如下：', indent=True)
    add_body_text(doc, '梯级1-2：上升/下降使能控制。上升键SET M1、上限位RST M1；下降键SET M2、下限位RST M2。', indent=True)
    add_body_text(doc, '梯级3-4：上升/下降输出，带互锁和急停保护。', indent=True)
    add_body_text(doc, '梯级5：停止复位。停止键→RST M1 + RST M2；急停键→ZRST M1 M2（区间复位，更高效）。ZRST指令可一次性复位M1到M2之间所有继电器，适合多状态同时清零的场景。', indent=True)
    add_body_text(doc, '梯级6：速度调节。INCP/DECP调整D0，比较指令控制速度输出Y011/Y012。', indent=True)
    add_body_text(doc, '梯级7-8：6流水灯控制。T20/T21组成0.167s脉冲发生器，6个灯依次轮流点亮，周期1s。6个流水灯是电动葫芦运行状态的重要视觉指示——在嘈杂的工业环境中，工人可能听不到电机运行声，但流水灯的光效能让他们直观看到设备正在运行中。', indent=True)
    add_body_text(doc, '梯级9：流水灯5/6 + 状态指示灯（上升绿灯、下降黄灯）。', indent=True)

    add_section_title(doc, '十、进度计划')
    progress_headers = ['时间', '设计内容']
    progress_rows = [
        ['6月22日', '介绍设计要求及任务，分析任务要求'],
        ['6月23日', '总体方案设计分析、确定所需电气元件'],
        ['6月24日', '输入/输出电路分析设计，流程图及程序设计'],
        ['6月25日', 'PLC程序编写'],
        ['6月26日', '系统完善，编写实训报告'],
    ]
    add_table(doc, progress_headers, progress_rows)

    add_section_title(doc, '十一、心得体会')
    add_body_text(doc, '本次电动葫芦控制系统的设计实训，让我全面理解了起重提升类设备的PLC控制特点和安全设计要求。以下是我在实训过程中的几点体会：', indent=True)
    add_body_text(doc, '第一，电动葫芦的安全设计是所有控制系统中最严格的，因为一旦发生故障，可能造成负载坠落伤人的严重后果。电磁制动器是电动葫芦的"最后一道防线"——它不需要任何电气控制信号，仅靠弹簧力就能在断电时自动抱闸。这种"失电安全"（Fail-Safe）设计理念值得在任何涉及重力负载的系统中采用。', indent=True)
    add_body_text(doc, '第二，限位开关在电动葫芦中扮演着"行程终点"的角色。上升限位防止吊钩冲顶（冲顶可能导致钢丝绳断裂、吊钩坠落），下降限位防止钢丝绳过度放出（可能导致钢丝绳从卷筒上完全脱落）。这两个限位开关必须绝对可靠，工业实践中通常采用"双限位"配置（两个串联的限位开关，任何一个触发即停止）。', indent=True)
    add_body_text(doc, '第三，ZRST区间复位指令的应用。在处理多状态复位场景（如急停时需要复位M1和M2）时，ZRST比多条RST指令更简洁、更可靠。ZRST M1 M2一次性地将M1到M2范围内所有继电器清零，避免了遗漏的风险。', indent=True)
    add_body_text(doc, '第四，6流水灯的脉冲分配设计。6个灯共享同一个脉冲源T20，每个灯在脉冲的不同相位点亮——这使得6个灯呈现出依次流动的视觉效果，且每个灯的切换时间精确到0.167秒。这种精确的时间控制正是PLC定时器的优势所在。', indent=True)
    add_body_text(doc, '第五，变频调速在电动葫芦中的应用具有实际意义。不同重量和尺寸的负载需要不同的提升速度——轻载可以快速提升以提高效率，重载需要慢速提升以保证稳定。通过PLC控制变频器的多段速功能，操作人员可以根据实际情况灵活选择提升速度，实现效率和安全的平衡。', indent=True)
    add_body_text(doc, '通过本次设计，我深刻理解了"安全是工业控制的灵魂"这句话的含义。任何一个控制系统的设计，首先要考虑的不是功能的丰富性，而是安全性和可靠性。只有保证了安全，才能谈功能和效率。', indent=True)

    path = os.path.join(OUT_DIR, '实训报告_电动葫芦控制系统.docx')
    doc.save(path)
    print(f'[OK] {path}')


if __name__ == '__main__':
    gen_concrete_mixer_report()
    gen_hub_tightener_report()
    gen_electric_hoist_report()
    print('\n全部实训报告生成完成!')
